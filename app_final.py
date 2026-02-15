"""
ScholarSync - Reference Generation Tool
A Streamlit web application for automated citation generation.
Supports DOI, arXiv, PMID, ISBN, URL, PDF, and manual entry.
"""

import streamlit as st
import requests
import re
import os
import io
import time
import json
import tempfile
import xml.etree.ElementTree as ET
from typing import Optional, Dict, Any, List, Tuple
from datetime import datetime
from collections import Counter
from urllib.parse import urlparse

import fitz  # PyMuPDF
from habanero import Crossref
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup

# ═══════════════════════════════════════════════════════════
# PAGE CONFIG & SESSION STATE
# ═══════════════════════════════════════════════════════════

st.set_page_config(
    page_title="ScholarSync",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded",
)

if "collection" not in st.session_state:
    st.session_state.collection = []
if "last_result" not in st.session_state:
    st.session_state.last_result = None
if "pending_duplicate" not in st.session_state:
    st.session_state.pending_duplicate = None

# ═══════════════════════════════════════════════════════════
# CONFIGURATION
# ═══════════════════════════════════════════════════════════

SUPPORTED_STYLES: Dict[str, str] = {
    "APA": "apa",
    "MLA": "modern-language-association",
    "Harvard": "harvard-cite-them-right",
    "Chicago": "chicago-author-date",
    "IEEE": "ieee",
    "Vancouver": "vancouver",
}

# Display names shown in the UI → internal style keys used in formatting logic
STYLE_DISPLAY_NAMES: Dict[str, str] = {
    "APA 7th Edition": "APA",
    "MLA 9th Edition": "MLA",
    "Harvard (Cite Them Right)": "Harvard",
    "Chicago 17th (Author-Date)": "Chicago",
    "IEEE": "IEEE",
    "Vancouver (ICMJE)": "Vancouver",
}

DOI_PATTERN = r"10\.\d{4,9}/[-._;()/:A-Z0-9]+[A-Z0-9]"
ARXIV_PATTERN = r"arXiv:\s*(\d{4}\.\d{4,5}(?:v\d+)?)"
PMID_PATTERN = r"PMID:\s*(\d{7,8})"
PMCID_PATTERN = r"PMC\s*(\d{6,7})"
ISBN_PATTERN = r"(?:ISBN[-: ]?)?(?:97[89][- ]?)?\d{1,5}[- ]?\d{1,7}[- ]?\d{1,7}[- ]?[\dX]"
URL_PATTERN = r"https?://(?:www\.)?[-a-zA-Z0-9@:%._\+~#=]{1,256}\.[a-zA-Z0-9()]{1,6}\b(?:[-a-zA-Z0-9()@:%_\+.~#?&/=]*)"

ARXIV_API_URL = "http://export.arxiv.org/api/query"
PUBMED_API_URL = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils"
GOOGLE_BOOKS_API_URL = "https://www.googleapis.com/books/v1/volumes"

MAX_RETRIES = 3
RETRY_DELAY = 1.0
BACKOFF_FACTOR = 2.0

# ═══════════════════════════════════════════════════════════
# VALIDATION & SANITIZATION
# ═══════════════════════════════════════════════════════════

def validate_doi(doi: str) -> bool:
    if not doi or not isinstance(doi, str):
        return False
    return bool(re.match(r"^10\.\d{4,9}/[-._;()/:A-Z0-9]+$", doi.strip(), re.IGNORECASE))

def validate_arxiv_id(arxiv_id: str) -> bool:
    if not arxiv_id or not isinstance(arxiv_id, str):
        return False
    return bool(re.match(r"^(arXiv:)?\d{4}\.\d{4,5}(v\d+)?$", arxiv_id.strip(), re.IGNORECASE))

def validate_pmid(pmid: str) -> bool:
    if not pmid or not isinstance(pmid, str):
        return False
    return bool(re.match(r"^\d{7,8}$", pmid.strip()))

def validate_pmcid(pmcid: str) -> bool:
    if not pmcid or not isinstance(pmcid, str):
        return False
    return bool(re.match(r"^PMC\d{6,7}$", pmcid.strip(), re.IGNORECASE))

def validate_isbn(isbn: str) -> bool:
    if not isbn or not isinstance(isbn, str):
        return False
    isbn = re.sub(r"[-\s]", "", isbn.strip())
    if len(isbn) == 10:
        if not re.match(r"^\d{9}[\dX]$", isbn, re.IGNORECASE):
            return False
        try:
            total = sum((10 - i) * (10 if c.upper() == "X" else int(c)) for i, c in enumerate(isbn))
            return total % 11 == 0
        except Exception:
            return False
    elif len(isbn) == 13:
        if not re.match(r"^(978|979)\d{10}$", isbn):
            return False
        try:
            total = sum((1 if i % 2 == 0 else 3) * int(c) for i, c in enumerate(isbn))
            return total % 10 == 0
        except Exception:
            return False
    return False

def validate_url(url: str) -> bool:
    if not url or not isinstance(url, str):
        return False
    try:
        result = urlparse(url.strip())
        return all([result.scheme in ["http", "https"], result.netloc])
    except Exception:
        return False

def validate_title(title: str) -> Tuple[bool, str]:
    if not title or not isinstance(title, str):
        return False, "Title cannot be empty"
    title = title.strip()
    if len(title) < 3:
        return False, "Title too short"
    if len(title) > 500:
        return False, "Title too long"
    if not re.search(r"[a-zA-Z]", title):
        return False, "Title must contain letters"
    return True, ""

def sanitize_doi(doi: str) -> str:
    if not doi:
        return doi
    doi = doi.strip()
    doi = re.sub(r"^(doi:|DOI:)\s*", "", doi, flags=re.IGNORECASE)
    doi = re.sub(r"^https?://doi\.org/", "", doi, flags=re.IGNORECASE)
    doi = re.sub(r"^https?://dx\.doi\.org/", "", doi, flags=re.IGNORECASE)
    return doi.strip()

def sanitize_arxiv_id(arxiv_id: str) -> str:
    if not arxiv_id:
        return arxiv_id
    return re.sub(r"^arXiv:\s*", "", arxiv_id.strip(), flags=re.IGNORECASE)

def sanitize_pmid(pmid: str) -> str:
    if not pmid:
        return pmid
    return re.sub(r"^PMID:\s*", "", pmid.strip(), flags=re.IGNORECASE)

def sanitize_pmcid(pmcid: str) -> str:
    if not pmcid:
        return pmcid
    pmcid = pmcid.strip()
    if not pmcid.upper().startswith("PMC"):
        pmcid = "PMC" + pmcid
    return pmcid.upper()

def sanitize_isbn(isbn: str) -> str:
    if not isbn:
        return isbn
    isbn = isbn.strip()
    isbn = re.sub(r"^ISBN[-:\s]*", "", isbn, flags=re.IGNORECASE)
    isbn = re.sub(r"[-\s]", "", isbn)
    return isbn.upper()

def sanitize_url(url: str) -> str:
    if not url:
        return url
    url = url.strip()
    if not url.startswith(("http://", "https://")):
        url = "https://" + url
    return url

# ═══════════════════════════════════════════════════════════
# RETRY LOGIC
# ═══════════════════════════════════════════════════════════

def retry_with_backoff(func, *args, max_attempts: int = MAX_RETRIES, **kwargs):
    delay = RETRY_DELAY
    for attempt in range(1, max_attempts + 1):
        try:
            return func(*args, **kwargs)
        except requests.RequestException:
            if attempt == max_attempts:
                return None
            time.sleep(delay)
            delay *= BACKOFF_FACTOR
        except Exception:
            return None
    return None

# ═══════════════════════════════════════════════════════════
# ML MODEL (CACHED)
# ═══════════════════════════════════════════════════════════

@st.cache_resource(show_spinner="Loading similarity model...")
def get_similarity_model():
    try:
        from sentence_transformers import SentenceTransformer
        return SentenceTransformer("all-MiniLM-L6-v2")
    except Exception:
        return "tfidf"

def calculate_similarity_advanced(text1: str, text2: str) -> float:
    try:
        model = get_similarity_model()
        if model == "tfidf":
            return _calculate_similarity_tfidf(text1, text2)
        from sklearn.metrics.pairwise import cosine_similarity
        embeddings = model.encode([text1, text2])
        return float(cosine_similarity([embeddings[0]], [embeddings[1]])[0][0])
    except Exception:
        return _calculate_similarity_tfidf(text1, text2)

def _calculate_similarity_tfidf(text1: str, text2: str) -> float:
    try:
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.metrics.pairwise import cosine_similarity
        tfidf = TfidfVectorizer()
        matrix = tfidf.fit_transform([text1, text2])
        return float(cosine_similarity(matrix[0:1], matrix[1:2])[0][0])
    except Exception:
        w1, w2 = set(text1.lower().split()), set(text2.lower().split())
        if not w1 or not w2:
            return 0.0
        return len(w1 & w2) / len(w1 | w2)

# ═══════════════════════════════════════════════════════════
# PDF PROCESSING
# ═══════════════════════════════════════════════════════════

def extract_identifiers_from_pdf_bytes(pdf_bytes: bytes) -> Dict[str, Optional[str]]:
    """Extract DOI, arXiv ID, and PMID from uploaded PDF bytes."""
    identifiers: Dict[str, Optional[str]] = {"doi": None, "arxiv": None, "pmid": None}
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = ""
        for page_num in range(min(2, len(doc))):
            text += doc[page_num].get_text()
        doc.close()

        if not text.strip():
            return identifiers

        # DOI
        doi_match = re.search(DOI_PATTERN, text, re.IGNORECASE)
        if doi_match:
            doi = sanitize_doi(doi_match.group(0))
            if validate_doi(doi):
                identifiers["doi"] = doi

        # arXiv
        arxiv_match = re.search(ARXIV_PATTERN, text, re.IGNORECASE)
        if arxiv_match:
            arxiv_id = sanitize_arxiv_id(arxiv_match.group(1))
            if validate_arxiv_id(arxiv_id):
                identifiers["arxiv"] = arxiv_id

        # PMID
        pmid_match = re.search(PMID_PATTERN, text, re.IGNORECASE)
        if pmid_match:
            pmid = sanitize_pmid(pmid_match.group(1))
            if validate_pmid(pmid):
                identifiers["pmid"] = pmid
    except Exception:
        pass
    return identifiers


def extract_title_from_pdf_bytes(pdf_bytes: bytes) -> Optional[str]:
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        if len(doc) == 0:
            return None
        text = doc[0].get_text()
        doc.close()
        if not text.strip():
            return None
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        skip_words = ["doi:", "http", "arxiv", "volume", "page", "journal",
                      "published", "copyright", "received", "accepted", "available", "pmid:", "pmc"]
        for line in lines[:15]:
            if any(s in line.lower() for s in skip_words):
                continue
            if 15 <= len(line) <= 300 and not line.isupper():
                return line
        if lines and len(lines[0]) >= 10:
            return lines[0]
        return None
    except Exception:
        return None


def extract_author_from_pdf_bytes(pdf_bytes: bytes) -> Optional[str]:
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        if len(doc) == 0:
            return None
        text = doc[0].get_text()
        doc.close()
        if not text.strip():
            return None
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        potential = []
        skip_words = ["abstract", "introduction", "keywords", "doi:", "http",
                      "journal", "volume", "university", "department", "email",
                      "copyright", "rights reserved", "published", "manuscript",
                      "corresponding author", "affiliation", "received", "accepted",
                      "article", "citation", "license", "creative commons", "arxiv", "pmid:", "pmc"]
        for i, line in enumerate(lines[:60]):
            if len(line) < 3 or len(line) > 500:
                continue
            if any(s in line.lower() for s in skip_words):
                continue
            if re.search(r"[A-Z][a-z]+", line):
                words = line.split()
                caps = [w for w in words if w and len(w) > 1 and w[0].isupper()]
                if 1 <= len(caps) <= 25:
                    has_sep = "," in line or " and " in line.lower() or "&" in line
                    looks_names = sum(1 for w in caps if len(w) >= 2 and sum(c.isalpha() for c in w) >= len(w) * 0.7) >= len(caps) * 0.5
                    if has_sep or looks_names:
                        score = len(caps) + (10 if has_sep else 0)
                        potential.append((i, line, score))
        if potential:
            potential.sort(key=lambda x: x[2], reverse=True)
            return " ".join(potential[0][1].split())
        return None
    except Exception:
        return None

# ═══════════════════════════════════════════════════════════
# API FETCHERS (CACHED)
# ═══════════════════════════════════════════════════════════

@st.cache_data(show_spinner=False, ttl=3600)
def get_paper_metadata(query: str) -> Optional[Dict[str, Any]]:
    """Fetch metadata from Crossref via DOI or title search."""
    query = query.strip()
    is_doi = query.startswith("10.")
    if is_doi:
        query = sanitize_doi(query)
        if not validate_doi(query):
            return None
    else:
        ok, _ = validate_title(query)
        if not ok:
            return None

    def fetch():
        cr = Crossref()
        if is_doi:
            result = cr.works(ids=query)
            return result["message"] if result and "message" in result else None
        else:
            results = cr.works(query=query, limit=1)
            if results and "message" in results and "items" in results["message"]:
                items = results["message"]["items"]
                return items[0] if items else None
            return None

    try:
        work = retry_with_backoff(fetch, max_attempts=MAX_RETRIES)
        if not work:
            return None
        all_authors = work.get("author", [])
        if all_authors:
            author_string = ", ".join(
                f"{a.get('given', '')} {a.get('family', '')}".strip() for a in all_authors
            )
        else:
            author_string = "Unknown Author"
        return {
            "title": work.get("title", ["Unknown Title"])[0] if work.get("title") else "Unknown Title",
            "author": author_string,
            "year": str(work.get("published", {}).get("date-parts", [[None]])[0][0] or "n.d."),
            "doi": work.get("DOI", "Unknown DOI"),
            "abstract": work.get("abstract", ""),
            "journal": work.get("container-title", [""])[0] if work.get("container-title") else "",
            "volume": str(work.get("volume", "")) if work.get("volume") else "",
            "issue": str(work.get("issue", "")) if work.get("issue") else "",
            "pages": work.get("page", ""),
        }
    except Exception:
        return None


@st.cache_data(show_spinner=False, ttl=3600)
def get_metadata_from_arxiv(arxiv_id: str) -> Optional[Dict[str, Any]]:
    arxiv_id = sanitize_arxiv_id(arxiv_id)
    if not validate_arxiv_id(arxiv_id):
        return None

    def fetch():
        resp = requests.get(ARXIV_API_URL, params={"id_list": arxiv_id, "max_results": 1}, timeout=10)
        if resp.status_code != 200:
            raise requests.RequestException(f"Status {resp.status_code}")
        root = ET.fromstring(resp.content)
        ns = {"atom": "http://www.w3.org/2005/Atom", "arxiv": "http://arxiv.org/schemas/atom"}
        entry = root.find("atom:entry", ns)
        if entry is None:
            return None
        title_el = entry.find("atom:title", ns)
        title = title_el.text.strip().replace("\n", " ") if title_el is not None else "Unknown Title"
        authors = []
        for a in entry.findall("atom:author", ns):
            n = a.find("atom:name", ns)
            if n is not None:
                authors.append(n.text.strip())
        pub_el = entry.find("atom:published", ns)
        year = pub_el.text[:4] if pub_el is not None else "n.d."
        doi_el = entry.find("arxiv:doi", ns)
        doi = doi_el.text.strip() if doi_el is not None else ""
        abs_el = entry.find("atom:summary", ns)
        abstract = abs_el.text.strip().replace("\n", " ") if abs_el is not None else ""
        return {"title": title, "author": ", ".join(authors) or "Unknown Author",
                "year": year, "doi": doi, "arxiv": arxiv_id, "abstract": abstract}

    try:
        return retry_with_backoff(fetch, max_attempts=MAX_RETRIES)
    except Exception:
        return None


@st.cache_data(show_spinner=False, ttl=3600)
def get_metadata_from_pubmed(pmid: str) -> Optional[Dict[str, Any]]:
    pmid = sanitize_pmid(pmid)
    if not validate_pmid(pmid):
        return None

    def fetch():
        resp = requests.get(f"{PUBMED_API_URL}/esummary.fcgi",
                            params={"db": "pubmed", "id": pmid, "retmode": "json"}, timeout=10)
        if resp.status_code != 200:
            raise requests.RequestException(f"Status {resp.status_code}")
        data = resp.json()
        if "result" not in data or pmid not in data["result"]:
            return None
        r = data["result"][pmid]
        title = r.get("title", "Unknown Title").strip()
        authors = [a.get("name", "") for a in r.get("authors", [])]
        year = r.get("pubdate", "n.d.")[:4] if "pubdate" in r else "n.d."
        doi = ""
        if "elocationid" in r:
            eloc = r["elocationid"]
            if eloc.startswith("doi:"):
                doi = eloc[4:].strip()
            elif "10." in eloc:
                doi = eloc.strip()
        if not doi and "articleids" in r:
            for aid in r["articleids"]:
                if aid.get("idtype") == "doi":
                    doi = aid.get("value", "")
                    break
        return {"title": title, "author": ", ".join(authors) or "Unknown Author",
                "year": year, "doi": doi, "pmid": pmid,
                "abstract": r.get("abstract", "")}

    try:
        return retry_with_backoff(fetch, max_attempts=MAX_RETRIES)
    except Exception:
        return None


@st.cache_data(show_spinner=False, ttl=3600)
def get_metadata_from_isbn(isbn: str) -> Optional[Dict[str, Any]]:
    isbn = sanitize_isbn(isbn)
    if not validate_isbn(isbn):
        return None

    def fetch():
        resp = requests.get(GOOGLE_BOOKS_API_URL, params={"q": f"isbn:{isbn}", "maxResults": 1}, timeout=10)
        if resp.status_code != 200:
            raise requests.RequestException(f"Status {resp.status_code}")
        data = resp.json()
        if "items" not in data or not data["items"]:
            return None
        info = data["items"][0]["volumeInfo"]
        pub_date = info.get("publishedDate", "n.d.")
        return {
            "title": info.get("title", "Unknown Title"),
            "author": ", ".join(info.get("authors", [])) or "Unknown Author",
            "year": pub_date[:4] if pub_date and len(pub_date) >= 4 else "n.d.",
            "publisher": info.get("publisher", ""),
            "isbn": isbn,
            "edition": info.get("printType", ""),
            "abstract": info.get("description", ""),
            "type": "book",
        }

    try:
        return retry_with_backoff(fetch, max_attempts=MAX_RETRIES)
    except Exception:
        return None


@st.cache_data(show_spinner=False, ttl=3600)
def get_metadata_from_url(url: str) -> Optional[Dict[str, Any]]:
    url = sanitize_url(url)
    if not validate_url(url):
        return None

    def fetch():
        headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
        resp = requests.get(url, headers=headers, timeout=10)
        if resp.status_code != 200:
            raise requests.RequestException(f"Status {resp.status_code}")
        soup = BeautifulSoup(resp.content, "html.parser")

        # Title
        title = None
        for getter in [
            lambda: (soup.find("meta", property="og:title") or {}).get("content"),
            lambda: (soup.find("meta", {"name": "twitter:title"}) or {}).get("content"),
            lambda: soup.title.string if soup.title else None,
            lambda: soup.find("h1").get_text().strip() if soup.find("h1") else None,
        ]:
            try:
                t = getter()
                if t:
                    title = t.strip()
                    break
            except Exception:
                continue
        title = title or "Unknown Title"

        # Author
        author = "Unknown Author"
        for getter in [
            lambda: (soup.find("meta", {"name": "author"}) or {}).get("content"),
            lambda: (soup.find("meta", property="article:author") or {}).get("content"),
            lambda: soup.find(["span", "div", "p"], class_=re.compile(r"author|byline", re.I)).get_text().strip(),
        ]:
            try:
                a = getter()
                if a and a.strip():
                    author = a.strip()
                    break
            except Exception:
                continue

        # Date
        year = "n.d."
        for getter in [
            lambda: (soup.find("meta", property="article:published_time") or {}).get("content"),
            lambda: (soup.find("meta", property="datePublished") or {}).get("content"),
            lambda: soup.find("time")["datetime"] if soup.find("time") else None,
        ]:
            try:
                d = getter()
                if d:
                    m = re.search(r"(\d{4})", d)
                    if m:
                        year = m.group(1)
                        break
            except Exception:
                continue

        # Publication name
        publication = "Unknown Publication"
        og_site = soup.find("meta", property="og:site_name")
        if og_site and og_site.get("content"):
            publication = og_site["content"].strip()
        else:
            publication = urlparse(url).netloc.replace("www.", "").split(".")[0].capitalize()

        # Description
        abstract = ""
        md = soup.find("meta", {"name": "description"})
        if md and md.get("content"):
            abstract = md["content"].strip()
        elif soup.find("meta", property="og:description"):
            abstract = soup.find("meta", property="og:description")["content"].strip()

        return {
            "title": title, "author": author, "year": year,
            "publication": publication, "url": url, "abstract": abstract,
            "access_date": datetime.now().strftime("%B %d, %Y"), "type": "web-article",
        }

    try:
        return retry_with_backoff(fetch, max_attempts=MAX_RETRIES)
    except Exception:
        return None


@st.cache_data(show_spinner=False, ttl=3600)
def get_metadata_from_pmcid(pmcid: str) -> Optional[Dict[str, Any]]:
    """Convert PMCID to PMID via NCBI ID converter, then fetch metadata."""
    pmcid = sanitize_pmcid(pmcid)
    if not validate_pmcid(pmcid):
        return None

    def fetch_pmid():
        resp = requests.get(
            "https://www.ncbi.nlm.nih.gov/pmc/utils/idconv/v1.0/",
            params={"ids": pmcid, "format": "json"},
            timeout=10,
        )
        if resp.status_code != 200:
            raise requests.RequestException(f"Status {resp.status_code}")
        data = resp.json()
        records = data.get("records", [])
        if records and "pmid" in records[0]:
            return records[0]["pmid"]
        return None

    try:
        pmid = retry_with_backoff(fetch_pmid, max_attempts=MAX_RETRIES)
        if pmid and validate_pmid(pmid):
            return get_metadata_from_pubmed(pmid)
        return None
    except Exception:
        return None


def get_metadata_with_fallback(identifiers: Dict[str, Optional[str]]) -> Optional[Dict[str, Any]]:
    if identifiers.get("doi"):
        meta = get_paper_metadata(identifiers["doi"])
        if meta:
            return meta
    if identifiers.get("arxiv"):
        meta = get_metadata_from_arxiv(identifiers["arxiv"])
        if meta:
            return meta
    if identifiers.get("pmid"):
        meta = get_metadata_from_pubmed(identifiers["pmid"])
        if meta:
            return meta
    return None

# ═══════════════════════════════════════════════════════════
# CITATION FORMATTING
# ═══════════════════════════════════════════════════════════

def format_author_names(author_string: str, style_name: str) -> str:
    if not author_string or author_string == "Unknown Author":
        return author_string
    separators = [", and ", " and ", ", ", " & "]
    authors = [author_string]
    for sep in separators:
        if sep in authors[0]:
            authors = [a.strip() for a in authors[0].split(sep) if a.strip()]
            break
    formatted = []
    for author in authors:
        author = re.sub(r"\d+", "", author).strip()
        parts = author.split()
        if not parts:
            continue
        if style_name == "APA":
            if len(parts) == 1:
                formatted.append(parts[0])
            else:
                initials = [f"{p[0]}." for p in parts[:-1] if p]
                formatted.append(f"{parts[-1]}, {' '.join(initials)}")
        elif style_name in ["MLA", "Chicago"]:
            if len(parts) == 1:
                formatted.append(parts[0])
            else:
                formatted.append(f"{parts[-1]}, {' '.join(parts[:-1])}")
        else:
            formatted.append(author)
    if not formatted:
        return author_string
    if len(formatted) == 1:
        return formatted[0]
    if len(formatted) == 2:
        return f"{formatted[0]}, & {formatted[1]}" if style_name == "APA" else f"{formatted[0]} and {formatted[1]}"
    if style_name == "APA":
        if len(formatted) <= 20:
            return ", ".join(formatted[:-1]) + f", & {formatted[-1]}"
        return ", ".join(formatted[:19]) + f", . . . {formatted[-1]}"
    return ", ".join(formatted[:-1]) + f", and {formatted[-1]}"


def get_parenthetical_citation(metadata: Dict[str, Any], style_name: str) -> str:
    try:
        author_full = metadata.get("author", "Unknown")
        if "," in author_full:
            last_name = author_full.split(",")[0].strip()
        else:
            parts = author_full.split()
            last_name = parts[-1] if parts else "Unknown"
        if "," in author_full or " and " in author_full.lower():
            last_name = last_name.split()[0] + " et al."
        year = metadata.get("year", "n.d.")
        if style_name in ["APA", "Harvard", "Chicago"]:
            return f"({last_name}, {year})"
        elif style_name == "MLA":
            return f"({last_name})"
        return f"({last_name}, {year})"
    except Exception:
        return "(Citation Error)"


def format_citation_from_metadata(metadata: Dict[str, Any], style_name: str) -> str:
    try:
        source_type = metadata.get("type", "journal")
        if source_type == "book":
            return _format_book_citation(metadata, style_name)
        elif source_type == "web-article":
            return _format_web_article_citation(metadata, style_name)
        return _format_journal_citation(metadata, style_name)
    except Exception:
        return f"{metadata.get('author', 'Unknown')} ({metadata.get('year', 'n.d.')}). {metadata.get('title', 'Unknown Title')}."


def _format_journal_citation(m: Dict[str, Any], style: str) -> str:
    fa = format_author_names(m.get("author", "Unknown Author"), style)
    t = m.get("title", "Unknown Title")
    y = m.get("year", "n.d.")
    doi = m.get("doi", "")
    j = m.get("journal", "")
    v = m.get("volume", "")
    iss = m.get("issue", "")
    p = m.get("pages", "")
    has_doi = doi and doi != "Unknown DOI" and validate_doi(doi)

    if style == "APA":
        c = f"{fa} ({y}). {t}."
        if j:
            c += f" {j}"
            if v:
                c += f", {v}"
                if iss:
                    c += f"({iss})"
            if p:
                c += f", {p}"
            c += "."
        if has_doi:
            c += f" https://doi.org/{doi}"
    elif style == "MLA":
        c = f'{fa}. "{t}."'
        if j:
            c += f" {j}"
            if v: c += f", vol. {v}"
            if iss: c += f", no. {iss}"
            c += f", {y}"
            if p: c += f", pp. {p}"
            c += "."
        else:
            c += f" {y}."
        if has_doi:
            c += f" doi:{doi}"
    elif style == "Harvard":
        c = f"{fa} ({y}) '{t}',"
        if j:
            c += f" {j}"
            if v:
                c += f", {v}"
                if iss: c += f"({iss})"
            if p: c += f", pp. {p}"
        c += "."
        if has_doi:
            c += f" doi: {doi}"
    elif style == "Chicago":
        c = f'{fa}. "{t}."'
        if j:
            c += f" {j}"
            if v: c += f" {v}"
            if iss: c += f", no. {iss}"
            c += f" ({y})"
            if p: c += f": {p}"
            c += "."
        else:
            c += f" {y}."
        if has_doi:
            c += f" https://doi.org/{doi}."
    elif style == "IEEE":
        c = f'{fa}, "{t},"'
        if j:
            c += f" {j}"
            if v: c += f", vol. {v}"
            if iss: c += f", no. {iss}"
            if p: c += f", pp. {p}"
        c += f", {y}."
        if has_doi:
            c += f" doi: {doi}."
    elif style == "Vancouver":
        c = f"{fa}. {t}."
        if j:
            c += f" {j}. {y}"
            if v:
                c += f";{v}"
                if iss: c += f"({iss})"
            if p: c += f":{p}"
            c += "."
        else:
            c += f" {y}."
        if has_doi:
            c += f" doi: {doi}"
    else:
        c = f"{fa} ({y}). {t}."
        if has_doi:
            c += f" https://doi.org/{doi}"
    return c


def _format_book_citation(m: Dict[str, Any], style: str) -> str:
    fa = format_author_names(m.get("author", "Unknown Author"), style)
    t = m.get("title", "Unknown Title")
    y = m.get("year", "n.d.")
    pub = m.get("publisher", "")
    ed = m.get("edition", "")
    if style == "APA":
        c = f"{fa} ({y}). {t}"
        if ed: c += f" ({ed})"
        c += "."
        if pub: c += f" {pub}."
    elif style == "MLA":
        c = f"{fa}. {t}."
        if ed: c += f" {ed},"
        if pub: c += f" {pub},"
        c += f" {y}."
    elif style == "Harvard":
        c = f"{fa} ({y}) {t}."
        if ed: c += f" {ed}."
        if pub: c += f" {pub}."
    elif style == "Chicago":
        c = f"{fa}. {y}. {t}."
        if ed: c += f" {ed}."
        if pub: c += f" {pub}."
    elif style == "IEEE":
        c = f"{fa}, {t}"
        if ed: c += f", {ed}"
        c += "."
        if pub: c += f" {pub},"
        c += f" {y}."
    elif style == "Vancouver":
        c = f"{fa}. {t}."
        if ed: c += f" {ed}."
        if pub: c += f" {pub};"
        c += f" {y}."
    else:
        c = f"{fa} ({y}). {t}."
        if pub: c += f" {pub}."
    return c


def _format_web_article_citation(m: Dict[str, Any], style: str) -> str:
    fa = format_author_names(m.get("author", "Unknown Author"), style)
    t = m.get("title", "Unknown Title")
    y = m.get("year", "n.d.")
    pub = m.get("publication", "")
    url = m.get("url", "")
    access = m.get("access_date", datetime.now().strftime("%B %d, %Y"))
    if style == "APA":
        c = f"{fa} ({y}). {t}."
        if pub: c += f" {pub}."
        if url: c += f" {url}"
    elif style == "MLA":
        c = f'{fa}. "{t}."'
        if pub: c += f" {pub},"
        c += f" {y},"
        if url: c += f" {url}."
        c += f" Accessed {access}."
    elif style == "Harvard":
        c = f"{fa} ({y}) '{t}'"
        if pub: c += f", {pub}"
        c += "."
        if url: c += f" Available at: {url} (Accessed: {access})."
    elif style == "Chicago":
        c = f'{fa}. "{t}."'
        if pub: c += f" {pub},"
        c += f" {y}."
        if url: c += f" {url}."
    elif style == "IEEE":
        c = f'{fa}, "{t},"'
        if pub: c += f" {pub},"
        c += f" {y}."
        if url: c += f" [Online]. Available: {url}"
    elif style == "Vancouver":
        c = f"{fa}. {t}."
        if pub: c += f" {pub}."
        c += f" {y}."
        if url: c += f" Available from: {url}"
    else:
        c = f"{fa} ({y}). {t}."
        if pub: c += f" {pub}."
        if url: c += f" {url}"
    return c

# ═══════════════════════════════════════════════════════════
# DUPLICATE DETECTION
# ═══════════════════════════════════════════════════════════

def check_duplicate_reference(
    new_ref: Dict[str, Any],
    existing_refs: List[Dict[str, Any]],
    title_threshold: float = 0.85,
    author_year_threshold: float = 0.70,
) -> Tuple[bool, Optional[Dict[str, Any]], Optional[str]]:
    if not existing_refs:
        return False, None, None
    nm = new_ref.get("meta", {})
    for er in existing_refs:
        em = er.get("meta", {})
        # Exact ID match
        for key in ["doi", "arxiv", "pmid"]:
            nv, ev = nm.get(key, "").strip(), em.get(key, "").strip()
            if nv and ev and nv != "Unknown DOI" and nv.lower() == ev.lower():
                return True, er, f"Identical {key.upper()}: {nv}"
        # Title similarity
        nt, et = nm.get("title", "").strip(), em.get("title", "").strip()
        if nt and et and nt != "Unknown Title":
            sim = calculate_similarity_advanced(nt, et)
            if sim >= title_threshold:
                return True, er, f"Similar title ({sim:.0%} match)"
        # Author + year
        na, ea = nm.get("author", "").lower(), em.get("author", "").lower()
        ny, ey = str(nm.get("year", "")), str(em.get("year", ""))
        if na and ea and ny and ey and ny == ey:
            if calculate_similarity_advanced(na, ea) >= author_year_threshold:
                return True, er, f"Same author & year"
    return False, None, None

# ═══════════════════════════════════════════════════════════
# AUTO-DETECTION & PIPELINE
# ═══════════════════════════════════════════════════════════

def _extract_identifier_from_url(url: str) -> Optional[Tuple[str, str]]:
    """Try to extract a known identifier (DOI, arXiv, PMID, PMCID, ISBN) from a URL.
    Returns (type, value) or None if the URL is not a recognised identifier link."""
    url_lower = url.lower()

    # DOI links: https://doi.org/10.xxxx/... or https://dx.doi.org/10.xxxx/...
    if "doi.org/" in url_lower:
        doi = sanitize_doi(url)
        if validate_doi(doi):
            return "doi", doi

    # Some publisher URLs embed the DOI in the path (e.g. /doi/10.xxxx/...)
    doi_in_path = re.search(r"/(10\.\d{4,9}/[^\s?#]+)", url)
    if doi_in_path:
        doi = doi_in_path.group(1).rstrip("/")
        if validate_doi(doi):
            return "doi", doi

    # arXiv links: https://arxiv.org/abs/2301.12345  or /pdf/2301.12345
    arxiv_match = re.search(r"arxiv\.org/(?:abs|pdf)/(\d{4}\.\d{4,5}(?:v\d+)?)", url_lower)
    if arxiv_match:
        aid = arxiv_match.group(1)
        if validate_arxiv_id(aid):
            return "arxiv", aid

    # PubMed links: https://pubmed.ncbi.nlm.nih.gov/12345678/
    pmid_match = re.search(r"pubmed\.ncbi\.nlm\.nih\.gov/(\d{7,8})", url_lower)
    if pmid_match:
        pid = pmid_match.group(1)
        if validate_pmid(pid):
            return "pmid", pid

    # PMC links: https://www.ncbi.nlm.nih.gov/pmc/articles/PMC1234567/
    pmc_match = re.search(r"ncbi\.nlm\.nih\.gov/pmc/articles/(PMC\d{6,7})", url, re.IGNORECASE)
    if pmc_match:
        pcid = pmc_match.group(1).upper()
        if validate_pmcid(pcid):
            return "pmcid", pcid

    # Google Books links containing ISBN
    isbn_match = re.search(r"[?&](?:isbn|vid=ISBN)[:=](\d{10,13})", url, re.IGNORECASE)
    if isbn_match:
        isbn = isbn_match.group(1)
        if validate_isbn(isbn):
            return "isbn", isbn

    return None


def detect_identifier_type(identifier_input: str) -> Tuple[str, str]:
    if not identifier_input or not isinstance(identifier_input, str):
        return "invalid", ""
    identifier_input = identifier_input.strip()
    # If the input looks like a URL, first try to extract a known identifier from it
    if identifier_input.startswith(("http://", "https://", "www.")):
        extracted = _extract_identifier_from_url(identifier_input)
        if extracted:
            return extracted
        # No known identifier found – treat as a generic URL
        url = sanitize_url(identifier_input)
        if validate_url(url):
            return "url", url
    # DOI
    if identifier_input.startswith("10.") or "doi.org/" in identifier_input.lower():
        doi = sanitize_doi(identifier_input)
        if validate_doi(doi):
            return "doi", doi
    # ISBN
    if "ISBN" in identifier_input.upper() or re.match(
        r"^(97[89])?\d{9}[\dX]$", identifier_input.replace("-", "").replace(" ", "")
    ):
        isbn = sanitize_isbn(identifier_input)
        if validate_isbn(isbn):
            return "isbn", isbn
    # arXiv
    if "arxiv" in identifier_input.lower() or re.match(r"^\d{4}\.\d{4,5}(v\d+)?$", identifier_input):
        aid = sanitize_arxiv_id(identifier_input)
        if validate_arxiv_id(aid):
            return "arxiv", aid
    # PMID
    if identifier_input.startswith("PMID:") or (identifier_input.isdigit() and 7 <= len(identifier_input) <= 8):
        pid = sanitize_pmid(identifier_input)
        if validate_pmid(pid):
            return "pmid", pid
    # PMCID
    if "PMC" in identifier_input.upper():
        pcid = sanitize_pmcid(identifier_input)
        if validate_pmcid(pcid):
            return "pmcid", pcid
    # Title fallback
    ok, _ = validate_title(identifier_input)
    if ok:
        return "title", identifier_input
    return "invalid", identifier_input


def process_identifier(identifier_input: str, style_name: str = "APA") -> Optional[Dict[str, Any]]:
    """Core pipeline: detect type -> fetch metadata -> format citation."""
    id_type, id_value = detect_identifier_type(identifier_input)
    if id_type == "invalid":
        return None

    metadata = None
    if id_type == "doi":
        metadata = get_paper_metadata(id_value)
    elif id_type == "arxiv":
        metadata = get_metadata_from_arxiv(id_value)
    elif id_type == "pmid":
        metadata = get_metadata_from_pubmed(id_value)
    elif id_type == "pmcid":
        metadata = get_metadata_from_pmcid(id_value)
    elif id_type == "isbn":
        metadata = get_metadata_from_isbn(id_value)
    elif id_type == "url":
        metadata = get_metadata_from_url(id_value)
    elif id_type == "title":
        metadata = get_paper_metadata(id_value)

    if not metadata:
        return None

    return {
        "meta": metadata,
        "full": format_citation_from_metadata(metadata, style_name),
        "parenthetical": get_parenthetical_citation(metadata, style_name),
        "style": style_name,
        "detected_type": id_type,
    }


def process_pdf(pdf_bytes: bytes, style_name: str = "APA") -> Optional[Dict[str, Any]]:
    """Process uploaded PDF bytes through the full pipeline."""
    identifiers = extract_identifiers_from_pdf_bytes(pdf_bytes)
    metadata = None
    if any(identifiers.values()):
        metadata = get_metadata_with_fallback(identifiers)
    if not metadata:
        title = extract_title_from_pdf_bytes(pdf_bytes)
        author = extract_author_from_pdf_bytes(pdf_bytes)
        if title:
            metadata = {
                "title": title,
                "author": author or "Unknown Author",
                "year": "n.d.",
                "doi": "",
            }
    if not metadata:
        return None
    return {
        "meta": metadata,
        "full": format_citation_from_metadata(metadata, style_name),
        "parenthetical": get_parenthetical_citation(metadata, style_name),
        "style": style_name,
        "detected_type": "pdf",
    }

# ═══════════════════════════════════════════════════════════
# WORD EXPORT (IN-MEMORY)
# ═══════════════════════════════════════════════════════════

def export_references_to_word(
    references: List[Dict[str, Any]],
    title: str = "References",
    sort_by_author: bool = True,
    include_abstracts: bool = True,
) -> Optional[bytes]:
    """Build a Word document in memory and return the bytes."""
    if not references:
        return None
    try:
        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        h = doc.add_heading(title, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        info = doc.add_paragraph()
        info.add_run(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n").italic = True
        info.add_run(f"Total References: {len(references)}\n").italic = True
        info.add_run(f"Citation Style: {references[0]['style']}\n").italic = True
        info.add_run("Generated by ScholarSync").italic = True
        doc.add_paragraph()
        doc.add_paragraph("─" * 80)
        doc.add_paragraph()

        refs = references.copy()
        if sort_by_author:
            def _sort_key(x):
                author = x["meta"].get("author", "")
                if not author or author == "Unknown Author":
                    return "zzz"
                # Get first author (before any comma separating multiple authors)
                first_author = author.split(",")[0].strip()
                # Last name is the last word of "FirstName LastName"
                parts = first_author.split()
                return parts[-1].lower() if parts else "zzz"
            refs.sort(key=_sort_key)

        # --- Section 1: Bibliography ---
        bib_h = doc.add_heading("Bibliography", level=2)
        bib_h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph()

        def _add_run(para, text, italic=False):
            if not text:
                return
            run = para.add_run(text)
            run.font.size = Pt(11)
            run.font.name = "Times New Roman"
            if italic:
                run.italic = True

        def _add_citation(para, citation_text, meta, style):
            citation_text = citation_text.replace("*", "")
            journal = meta.get("journal", "")
            volume = meta.get("volume", "")
            if journal and journal in citation_text:
                idx = citation_text.find(journal)
                _add_run(para, citation_text[:idx])
                _add_run(para, journal, italic=True)
                after = citation_text[idx + len(journal):]
                if volume and style in ["APA", "Harvard"]:
                    vol_pat = rf"(,\s*)({re.escape(volume)})(\(?)"
                    vm = re.search(vol_pat, after)
                    if vm:
                        _add_run(para, after[:vm.start()])
                        _add_run(para, vm.group(1))
                        _add_run(para, vm.group(2), italic=True)
                        _add_run(para, after[vm.end() - len(vm.group(3)):])
                    else:
                        _add_run(para, after)
                else:
                    _add_run(para, after)
            else:
                _add_run(para, citation_text)

        for ref in refs:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.space_after = Pt(12)
            _add_citation(p, ref["full"], ref.get("meta", {}), ref.get("style", "APA"))

        doc.add_page_break()

        # --- Section 2: In-Text Citations ---
        it_h = doc.add_heading("In-Text Citations", level=2)
        it_h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph()
        exp = doc.add_paragraph()
        exp_run = exp.add_run("Use these in-text citations when referencing the papers in your work:")
        exp_run.font.size = Pt(10)
        exp_run.italic = True
        doc.add_paragraph()

        for i, ref in enumerate(refs, 1):
            itp = doc.add_paragraph()
            nr = itp.add_run(f"[{i}] ")
            nr.bold = True
            nr.font.size = Pt(11)
            nr.font.name = "Times New Roman"
            cr = itp.add_run(ref["parenthetical"])
            cr.font.size = Pt(11)
            cr.font.name = "Times New Roman"
            tp = doc.add_paragraph()
            tp.paragraph_format.left_indent = Inches(0.3)
            ttxt = f"→ {ref['meta']['title'][:80]}{'...' if len(ref['meta']['title']) > 80 else ''}"
            tr = tp.add_run(ttxt)
            tr.font.size = Pt(9)
            tr.italic = True
            tp.paragraph_format.space_after = Pt(12)

        # --- Section 3: Abstracts ---
        if include_abstracts:
            abs_refs = [r for r in refs if r["meta"].get("abstract")]
            if abs_refs:
                doc.add_page_break()
                ah = doc.add_heading("Abstracts", level=2)
                ah.alignment = WD_ALIGN_PARAGRAPH.LEFT
                doc.add_paragraph()
                ae = doc.add_paragraph()
                ae_run = ae.add_run("Abstracts extracted from paper metadata (where available):")
                ae_run.font.size = Pt(10)
                ae_run.italic = True
                doc.add_paragraph()
                for idx, ref in enumerate(abs_refs, 1):
                    hp = doc.add_paragraph()
                    hnr = hp.add_run(f"[{idx}] ")
                    hnr.bold = True; hnr.font.size = Pt(11); hnr.font.name = "Times New Roman"
                    hcr = hp.add_run(f"{ref['meta']['author']} ({ref['meta']['year']})")
                    hcr.font.size = Pt(11); hcr.font.name = "Times New Roman"; hcr.bold = True
                    ttp = doc.add_paragraph()
                    ttp.paragraph_format.left_indent = Inches(0.3)
                    ttr = ttp.add_run(ref["meta"]["title"])
                    ttr.font.size = Pt(10); ttr.italic = True; ttr.font.name = "Times New Roman"
                    abp = doc.add_paragraph()
                    abp.paragraph_format.left_indent = Inches(0.3)
                    abp.paragraph_format.right_indent = Inches(0.3)
                    abp.paragraph_format.space_after = Pt(16)
                    abs_text = ref["meta"]["abstract"]
                    if len(abs_text) > 1500:
                        abs_text = abs_text[:1500] + "... [truncated]"
                    abr = abp.add_run(abs_text)
                    abr.font.size = Pt(10); abr.font.name = "Times New Roman"
                    doc.add_paragraph("─" * 60)
                    doc.add_paragraph()

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception:
        return None

# ═══════════════════════════════════════════════════════════
# ABSTRACT PROCESSING
# ═══════════════════════════════════════════════════════════

def process_abstract(abstract: str) -> Dict[str, Any]:
    wc = len(abstract.split()) if abstract else 0
    display = (" ".join(abstract.split()[:150]) + "...") if wc > 150 else abstract
    return {"full": abstract, "display": display, "word_count": wc}

# ═══════════════════════════════════════════════════════════
# HELPER: DISPLAY RESULT CARD
# ═══════════════════════════════════════════════════════════

def _show_result(result: Dict[str, Any], key_prefix: str = ""):
    """Show a citation result card with add-to-collection button."""
    meta = result["meta"]
    detected = result.get("detected_type", "unknown").upper()

    st.markdown(f"**Detected type:** `{detected}`")

    col1, col2 = st.columns([3, 1])
    with col1:
        st.markdown(f"**Title:** {meta.get('title', 'N/A')}")
        st.markdown(f"**Author:** {meta.get('author', 'N/A')}")
        st.markdown(f"**Year:** {meta.get('year', 'N/A')}")
        for field, label in [("doi", "DOI"), ("arxiv", "arXiv"), ("pmid", "PMID"),
                             ("isbn", "ISBN"), ("publisher", "Publisher"),
                             ("journal", "Journal"), ("url", "URL")]:
            val = meta.get(field)
            if val:
                st.markdown(f"**{label}:** {val}")
    with col2:
        vol = meta.get("volume", "")
        iss = meta.get("issue", "")
        pgs = meta.get("pages", "")
        if vol:
            st.markdown(f"**Volume:** {vol}")
        if iss:
            st.markdown(f"**Issue:** {iss}")
        if pgs:
            st.markdown(f"**Pages:** {pgs}")

    # Abstract
    if meta.get("abstract"):
        with st.expander("📄 Abstract", expanded=False):
            ab = process_abstract(meta["abstract"])
            st.write(ab["full"])
            st.caption(f"{ab['word_count']} words")

    # Citations
    st.markdown("---")
    st.markdown("**Full Citation:**")
    st.code(result["full"], language=None)
    st.markdown("**In-Text Citation:**")
    st.code(result["parenthetical"], language=None)

    # Add button
    btn_key = f"add_{key_prefix}_{meta.get('title', '')[:30]}"
    if st.button("➕ Add to Collection", key=btn_key, type="primary"):
        if st.session_state.get("duplicate_detection", True):
            is_dup, dup_ref, reason = check_duplicate_reference(result, st.session_state.collection)
            if is_dup:
                st.session_state.pending_duplicate = {"result": result, "reason": reason, "dup": dup_ref, "key": btn_key}
                st.rerun()
        else:
            is_dup = False
        if not st.session_state.get("pending_duplicate"):
            st.session_state.collection.append(result)
            st.toast("✅ Added to collection!")
            st.rerun()

# ═══════════════════════════════════════════════════════════
# SIDEBAR
# ═══════════════════════════════════════════════════════════

with st.sidebar:
    st.image("https://img.icons8.com/fluency/96/book-shelf.png", width=60)
    st.title("ScholarSync")
    st.caption("Smart Reference Generator")
    st.divider()

    style_display = st.selectbox("📝 Citation Style", list(STYLE_DISPLAY_NAMES.keys()), index=0)
    style = STYLE_DISPLAY_NAMES[style_display]
    st.divider()

    st.metric("📚 References Collected", len(st.session_state.collection))

    if st.session_state.collection:
        word_bytes = export_references_to_word(st.session_state.collection, sort_by_author=True)
        if word_bytes:
            st.download_button(
                "⬇️ Download Word (.docx)",
                data=word_bytes,
                file_name=f"ScholarSync_References_{datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

        if st.button("🗑️ Clear All References", use_container_width=True):
            st.session_state.collection.clear()
            st.session_state.last_result = None
            st.session_state.pending_duplicate = None
            st.rerun()

    st.divider()
    st.toggle("🔍 Duplicate Detection", value=True, key="duplicate_detection",
             help="Uses an ML similarity model to detect duplicate references before adding.")
    st.divider()
    st.caption("Supports: DOI · arXiv · PMID · ISBN · URL · PDF · Title")

# ═══════════════════════════════════════════════════════════
# DUPLICATE WARNING MODAL
# ═══════════════════════════════════════════════════════════

if st.session_state.pending_duplicate:
    dup_info = st.session_state.pending_duplicate
    st.warning(f"⚠️ **Duplicate detected:** {dup_info['reason']}")
    dup_meta = dup_info["dup"]["meta"]
    st.caption(f"Existing: *{dup_meta.get('title', 'N/A')[:80]}* — {dup_meta.get('author', 'N/A')[:50]} ({dup_meta.get('year', 'N/A')})")
    c1, c2 = st.columns(2)
    with c1:
        if st.button("Add Anyway", type="primary"):
            st.session_state.collection.append(dup_info["result"])
            st.session_state.pending_duplicate = None
            st.toast("✅ Added (duplicate override)")
            st.rerun()
    with c2:
        if st.button("Cancel"):
            st.session_state.pending_duplicate = None
            st.rerun()
    st.divider()

# ═══════════════════════════════════════════════════════════
# MAIN AREA
# ═══════════════════════════════════════════════════════════

st.title("📚 ScholarSync")
st.caption("Paste a DOI, ISBN, arXiv ID, PMID, URL, title, or upload a PDF to generate citations instantly.")

tab_search, tab_pdf, tab_manual, tab_collection = st.tabs(
    ["🔍 Universal Search", "📄 PDF Upload", "✍️ Manual Entry", "📚 My Collection"]
)

# ── Tab 1: Universal Search ──────────────────────────────

with tab_search:
    identifier = st.text_input(
        "Paste a DOI, ISBN, arXiv ID, PMID, URL, or Title",
        placeholder="e.g. 10.1038/nature12345  or  https://www.bbc.com/news/article  or  978-0-7475-3269-9",
        key="search_input",
    )
    if st.button("Generate Citation", key="search_go", type="primary", use_container_width=True):
        if not identifier.strip():
            st.warning("Please enter an identifier or title.")
        else:
            with st.spinner("Fetching metadata..."):
                result = process_identifier(identifier.strip(), style)
            if result:
                st.session_state.last_result = result
            else:
                st.error("❌ Could not retrieve metadata. Check the identifier and try again.")

    if st.session_state.last_result and st.session_state.last_result.get("detected_type") != "pdf_tab":
        _show_result(st.session_state.last_result, key_prefix="search")

# ── Tab 2: PDF Upload ─────────────────────────────────────

with tab_pdf:
    uploaded = st.file_uploader("Upload a research paper PDF", type=["pdf"], key="pdf_upload")
    if uploaded is not None:
        pdf_bytes = uploaded.read()
        with st.spinner("Extracting identifiers from PDF..."):
            result = process_pdf(pdf_bytes, style)
        if result:
            result["detected_type"] = "pdf_tab"
            _show_result(result, key_prefix="pdf")
        else:
            st.error("❌ Could not extract enough information from the PDF.")

# ── Tab 3: Manual Entry ───────────────────────────────────

with tab_manual:
    m_source_type = st.radio(
        "Source Type",
        ["Journal Article", "Book", "Web / News Article"],
        horizontal=True,
        key="manual_source_type",
    )
    with st.form("manual_form"):
        m_title = st.text_input("Title *", placeholder="The Impact of Machine Learning on Healthcare")
        m_author = st.text_input("Author(s) *", placeholder="John Smith, Jane Doe")
        mc1, mc2 = st.columns(2)
        with mc1:
            m_year = st.text_input("Year", value="n.d.", placeholder="2023")
        with mc2:
            if m_source_type == "Journal Article":
                m_doi = st.text_input("DOI (optional)", placeholder="10.1234/example")
            elif m_source_type == "Book":
                m_isbn = st.text_input("ISBN (optional)", placeholder="978-0-7475-3269-9")
            else:
                m_url = st.text_input("URL *", placeholder="https://www.bbc.com/news/article")

        if m_source_type == "Journal Article":
            mc3, mc4 = st.columns(2)
            with mc3:
                m_journal = st.text_input("Journal (optional)", placeholder="Nature")
            with mc4:
                m_volume = st.text_input("Volume (optional)", placeholder="46")
            mc5, mc6 = st.columns(2)
            with mc5:
                m_issue = st.text_input("Issue (optional)", placeholder="2")
            with mc6:
                m_pages = st.text_input("Pages (optional)", placeholder="123-145")
        elif m_source_type == "Book":
            mc3, mc4 = st.columns(2)
            with mc3:
                m_publisher = st.text_input("Publisher (optional)", placeholder="Oxford University Press")
            with mc4:
                m_edition = st.text_input("Edition (optional)", placeholder="3rd")
        else:  # Web / News Article
            mc3, mc4 = st.columns(2)
            with mc3:
                m_publication = st.text_input("Publication / Website Name", placeholder="BBC News")
            with mc4:
                m_access_date = st.text_input("Access Date", value=datetime.now().strftime("%B %d, %Y"))

        submitted = st.form_submit_button("Generate Citation", type="primary", use_container_width=True)
        if submitted:
            if not m_title.strip() or not m_author.strip():
                st.warning("Title and Author are required.")
            elif m_source_type == "Web / News Article" and not m_url.strip():
                st.warning("URL is required for web/news articles.")
            else:
                if m_source_type == "Journal Article":
                    meta = {
                        "title": m_title.strip(),
                        "author": m_author.strip(),
                        "year": m_year.strip() or "n.d.",
                        "doi": m_doi.strip(),
                        "journal": m_journal.strip(),
                        "volume": m_volume.strip(),
                        "issue": m_issue.strip(),
                        "pages": m_pages.strip(),
                    }
                elif m_source_type == "Book":
                    meta = {
                        "title": m_title.strip(),
                        "author": m_author.strip(),
                        "year": m_year.strip() or "n.d.",
                        "publisher": m_publisher.strip(),
                        "edition": m_edition.strip(),
                        "isbn": m_isbn.strip() if m_isbn.strip() else "",
                        "type": "book",
                    }
                else:  # Web / News Article
                    meta = {
                        "title": m_title.strip(),
                        "author": m_author.strip(),
                        "year": m_year.strip() or "n.d.",
                        "publication": m_publication.strip(),
                        "url": m_url.strip(),
                        "access_date": m_access_date.strip(),
                        "type": "web-article",
                    }
                st.session_state.manual_result = {
                    "meta": meta,
                    "full": format_citation_from_metadata(meta, style),
                    "parenthetical": get_parenthetical_citation(meta, style),
                    "style": style,
                    "detected_type": "manual",
                }

    # Display result outside the form so st.button works
    if st.session_state.get("manual_result"):
        _show_result(st.session_state.manual_result, key_prefix="manual")

# ── Tab 4: Collection & Analytics ─────────────────────────

with tab_collection:
    refs = st.session_state.collection
    if not refs:
        st.info("No references collected yet. Use the other tabs to generate and add citations.")
    else:
        st.subheader(f"📚 {len(refs)} Reference(s)")

        # Year analytics
        years = []
        for ref in refs:
            ys = str(ref["meta"].get("year", "")).strip()
            if ys.isdigit() and 1900 <= int(ys) <= 2100:
                years.append(int(ys))

        if years:
            counts = Counter(years)
            chart_data = {str(y): counts[y] for y in sorted(counts)}
            st.bar_chart(chart_data)
            scol1, scol2, scol3, scol4 = st.columns(4)
            scol1.metric("Total", len(years))
            scol2.metric("Earliest", min(years))
            scol3.metric("Latest", max(years))
            scol4.metric("Range", f"{max(years) - min(years)} yrs")

        st.divider()

        # List each reference
        for i, ref in enumerate(refs):
            with st.expander(f"**[{i+1}]** {ref['meta'].get('title', 'N/A')[:80]}", expanded=False):
                st.markdown(f"**Author:** {ref['meta'].get('author', 'N/A')}")
                st.markdown(f"**Year:** {ref['meta'].get('year', 'N/A')}")
                st.code(ref["full"], language=None)
                st.code(ref["parenthetical"], language=None)
                if ref["meta"].get("abstract"):
                    ab = process_abstract(ref["meta"]["abstract"])
                    st.caption(f"Abstract ({ab['word_count']} words): {ab['display'][:200]}...")
                if st.button("🗑️ Remove", key=f"remove_{i}"):
                    st.session_state.collection.pop(i)
                    st.rerun()
