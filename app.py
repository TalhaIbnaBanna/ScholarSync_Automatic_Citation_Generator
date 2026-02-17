"""
ScholarSync - Reference Generation Tool
A Streamlit web application for automated citation generation.
Supports DOI, arXiv, PMID, ISBN, URL, PDF, and manual entry.
"""

# Standard library imports
import streamlit as st
import requests
import re
import os
import io
import time
import json
import html
import tempfile
import xml.etree.ElementTree as ET
from typing import Optional, Dict, Any, List, Tuple
from datetime import datetime
from collections import Counter
from urllib.parse import urlparse

# Third-party imports
import fitz  # PyMuPDF – used for PDF text extraction
from habanero import Crossref  # Python client for the Crossref REST API (DOI lookups)
from docx import Document  # python-docx – generates Word (.docx) files for export
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup  # HTML parser for web-scraping article metadata from URLs
import streamlit.components.v1 as components  # For custom HTML components with JS

# ═══════════════════════════════════════════════════════════
# PAGE CONFIG & SESSION STATE
# ═══════════════════════════════════════════════════════════

st.set_page_config(
    page_title="ScholarSync",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Initialize session state variables that persist across Streamlit reruns:
# - collection: list of saved reference dicts (the user's bibliography)
# - last_result: most recent search result displayed in the Search tab
# - pending_duplicate: holds duplicate warning info when a conflict is detected
if "collection" not in st.session_state:
    st.session_state.collection = []
if "last_result" not in st.session_state:
    st.session_state.last_result = None
if "pending_duplicate" not in st.session_state:
    st.session_state.pending_duplicate = None

# ═══════════════════════════════════════════════════════════
# CONFIGURATION
# ═══════════════════════════════════════════════════════════

# Mapping of internal style keys to their Crossref/CSL style identifiers
SUPPORTED_STYLES: Dict[str, str] = {
    "APA": "apa",
    "MLA": "modern-language-association",
    "Harvard": "harvard-cite-them-right",
    "Chicago": "chicago-author-date",
    "IEEE": "ieee",
    "Vancouver": "vancouver",
}

# Display names shown in the UI → internal style keys used in formatting logic
STYLE_DISPLAY_NAMES: Dict[str, str] = {
    "APA 7th Edition": "APA",
    "MLA 9th Edition": "MLA",
    "Harvard (Cite Them Right)": "Harvard",
    "Chicago 17th (Author-Date)": "Chicago",
    "IEEE": "IEEE",
    "Vancouver (ICMJE)": "Vancouver",
}

# Regex patterns for detecting identifier types in user input and PDF text
DOI_PATTERN = r"10\.\d{4,9}/[-._;()/:A-Z0-9]+[A-Z0-9]"
ARXIV_PATTERN = r"arXiv:\s*(\d{4}\.\d{4,5}(?:v\d+)?)"
PMID_PATTERN = r"PMID:\s*(\d{1,9})"
PMCID_PATTERN = r"PMC\s*(\d{1,9})"
ISBN_PATTERN = r"(?:ISBN[-: ]?)?(?:97[89][- ]?)?\d{1,5}[- ]?\d{1,7}[- ]?\d{1,7}[- ]?[\dX]"
URL_PATTERN = r"https?://(?:www\.)?[-a-zA-Z0-9@:%._\+~#=]{1,256}\.[a-zA-Z0-9()]{1,6}\b(?:[-a-zA-Z0-9()@:%_\+.~#?&/=]*)"

# External API base URLs
ARXIV_API_URL = "http://export.arxiv.org/api/query"
PUBMED_API_URL = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils"
GOOGLE_BOOKS_API_URL = "https://www.googleapis.com/books/v1/volumes"

# Retry configuration for API calls (exponential backoff)
MAX_RETRIES = 3
RETRY_DELAY = 1.0       # Initial delay in seconds
BACKOFF_FACTOR = 2.0    # Multiplier applied to delay after each failed attempt

# ═══════════════════════════════════════════════════════════
# VALIDATION & SANITIZATION
# ═══════════════════════════════════════════════════════════

def validate_doi(doi: str) -> bool:
    """Check if a string is a valid DOI (Digital Object Identifier).
    DOIs always start with '10.' followed by a registrant code and a suffix separated by '/'.
    Example valid DOI: 10.1038/nature12345"""
    if not doi or not isinstance(doi, str):
        return False
    # Regex: starts with 10., then 4-9 digit registrant code, then /, then suffix characters
    # Suffix allows letters, digits, and common DOI punctuation including <>, #, ~, [], etc.
    return bool(re.match(r"^10\.\d{4,9}/\S+$", doi.strip(), re.IGNORECASE))

def validate_arxiv_id(arxiv_id: str) -> bool:
    """Check if a string is a valid arXiv preprint ID.
    Format: YYMM.NNNNN with optional version suffix (e.g., 2301.12345v2).
    The 'arXiv:' prefix is optional."""
    if not arxiv_id or not isinstance(arxiv_id, str):
        return False
    # Regex: optional 'arXiv:' prefix, then YYMM.NNNNN, optional version vN
    return bool(re.match(r"^(arXiv:)?\d{4}\.\d{4,5}(v\d+)?$", arxiv_id.strip(), re.IGNORECASE))

def validate_pmid(pmid: str) -> bool:
    """Check if a string is a valid PubMed ID (PMID).
    PMIDs are purely numeric identifiers assigned to biomedical articles.
    Example: 39437913"""
    if not pmid or not isinstance(pmid, str):
        return False
    # PMIDs are 1-9 digit numbers (older articles have shorter IDs)
    return bool(re.match(r"^\d{1,9}$", pmid.strip()))

def validate_pmcid(pmcid: str) -> bool:
    """Check if a string is a valid PubMed Central ID (PMCID).
    PMCIDs have the format 'PMC' followed by digits (e.g., PMC11312261).
    These identify freely accessible full-text articles in the PMC archive."""
    if not pmcid or not isinstance(pmcid, str):
        return False
    # 'PMC' prefix followed by 1-9 digits (newer articles can have 8+ digits)
    return bool(re.match(r"^PMC\d{1,9}$", pmcid.strip(), re.IGNORECASE))

def validate_isbn(isbn: str) -> bool:
    """Validate ISBN-10 (mod-11 checksum) or ISBN-13 (mod-10 checksum)."""
    if not isbn or not isinstance(isbn, str):
        return False
    isbn = re.sub(r"[-\s]", "", isbn.strip())
    if len(isbn) == 10:
        # ISBN-10: 9 digits + check digit (0-9 or X); weighted sum must be divisible by 11
        if not re.match(r"^\d{9}[\dX]$", isbn, re.IGNORECASE):
            return False
        try:
            total = sum((10 - i) * (10 if c.upper() == "X" else int(c)) for i, c in enumerate(isbn))
            return total % 11 == 0
        except Exception:
            return False
    elif len(isbn) == 13:
        # ISBN-13: must start with 978 or 979; alternating weights of 1 and 3, sum mod 10 == 0
        if not re.match(r"^(978|979)\d{10}$", isbn):
            return False
        try:
            total = sum((1 if i % 2 == 0 else 3) * int(c) for i, c in enumerate(isbn))
            return total % 10 == 0
        except Exception:
            return False
    return False

def validate_url(url: str) -> bool:
    """Check if a string is a valid HTTP/HTTPS URL using Python's urlparse.
    Requires both a valid scheme (http/https) and a network location (domain)."""
    if not url or not isinstance(url, str):
        return False
    try:
        result = urlparse(url.strip())
        # Must have http/https scheme AND a domain (netloc)
        return all([result.scheme in ["http", "https"], result.netloc])
    except Exception:
        return False

def validate_title(title: str) -> Tuple[bool, str]:
    """Validate a freeform title string for use as a Crossref search query.
    Returns (is_valid, error_message). Used as a last-resort identifier type
    when the input doesn't match any structured identifier format."""
    if not title or not isinstance(title, str):
        return False, "Title cannot be empty"
    title = title.strip()
    if len(title) < 3:
        return False, "Title too short"
    if len(title) > 500:
        return False, "Title too long"
    if not re.search(r"[a-zA-Z]", title):
        return False, "Title must contain letters"
    return True, ""

def sanitize_doi(doi: str) -> str:
    """Strip common prefixes from DOI input to get the bare identifier.
    Handles 'doi:10.xxxx', 'https://doi.org/10.xxxx', 'https://dx.doi.org/10.xxxx'."""
    if not doi:
        return doi
    doi = doi.strip()
    doi = re.sub(r"^(doi:|DOI:)\s*", "", doi, flags=re.IGNORECASE)   # Remove 'doi:' prefix
    doi = re.sub(r"^https?://doi\.org/", "", doi, flags=re.IGNORECASE)   # Remove doi.org URL
    doi = re.sub(r"^https?://dx\.doi\.org/", "", doi, flags=re.IGNORECASE)  # Remove dx.doi.org URL
    return doi.strip()

def sanitize_arxiv_id(arxiv_id: str) -> str:
    """Remove the 'arXiv:' prefix if present, leaving just the numeric ID."""
    if not arxiv_id:
        return arxiv_id
    return re.sub(r"^arXiv:\s*", "", arxiv_id.strip(), flags=re.IGNORECASE)

def sanitize_pmid(pmid: str) -> str:
    """Remove the 'PMID:' prefix if present, leaving just the numeric ID."""
    if not pmid:
        return pmid
    return re.sub(r"^PMID:\s*", "", pmid.strip(), flags=re.IGNORECASE)

def sanitize_pmcid(pmcid: str) -> str:
    if not pmcid:
        return pmcid
    pmcid = pmcid.strip()
    # Extract the numeric part from various input formats:
    # "PMC1234567", "PMCID: 1234567", "PMCID1234567", "PMC 1234567"
    m = re.search(r"(\d{1,9})", pmcid)
    if m:
        return "PMC" + m.group(1)
    return pmcid.upper()

def sanitize_isbn(isbn: str) -> str:
    """Clean an ISBN string by removing the 'ISBN' prefix, hyphens, and spaces.
    Example: 'ISBN 978-0-7475-3269-9' → '9780747532699'"""
    if not isbn:
        return isbn
    isbn = isbn.strip()
    isbn = re.sub(r"^ISBN[-:\s]*", "", isbn, flags=re.IGNORECASE)  # Remove 'ISBN' prefix
    isbn = re.sub(r"[-\s]", "", isbn)  # Remove hyphens and spaces
    return isbn.upper()

def sanitize_url(url: str) -> str:
    """Ensure URL has a proper scheme. Adds 'https://' if missing."""
    if not url:
        return url
    url = url.strip()
    if not url.startswith(("http://", "https://")):
        url = "https://" + url
    return url

def clean_abstract(abstract: str) -> str:
    """Strip JATS XML, HTML tags, and other markup from abstract text.
    Crossref often returns abstracts wrapped in JATS tags like <jats:p>, <jats:sec>, etc."""
    if not abstract:
        return ""
    # Remove all XML/HTML tags (handles namespaced tags like <jats:p> too)
    cleaned = re.sub(r"<[^>]+>", " ", abstract)
    # Collapse multiple whitespace into single spaces
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return cleaned

# ═══════════════════════════════════════════════════════════
# RETRY LOGIC
# ═══════════════════════════════════════════════════════════

def retry_with_backoff(func, *args, max_attempts: int = MAX_RETRIES, **kwargs):
    """Execute `func` with exponential backoff on network errors.
    Returns None if all attempts fail or a non-network exception occurs."""
    delay = RETRY_DELAY
    for attempt in range(1, max_attempts + 1):
        try:
            return func(*args, **kwargs)
        except requests.RequestException:
            # Network/HTTP error — retry after increasing delay
            if attempt == max_attempts:
                return None
            time.sleep(delay)
            delay *= BACKOFF_FACTOR
        except Exception:
            # Non-network error (e.g. JSON parse failure) — don't retry
            return None
    return None

# ═══════════════════════════════════════════════════════════
# ML MODEL (CACHED)
# ═══════════════════════════════════════════════════════════

@st.cache_resource(show_spinner="Loading similarity model...")
def get_similarity_model():
    """Load the sentence-transformer model for semantic duplicate detection.
    Falls back to TF-IDF string similarity if the model can't be loaded."""
    try:
        from sentence_transformers import SentenceTransformer
        return SentenceTransformer("all-MiniLM-L6-v2")
    except Exception:
        return "tfidf"  # Sentinel value indicating fallback mode

def calculate_similarity_advanced(text1: str, text2: str) -> float:
    """Compute semantic similarity between two text strings using the ML model.
    Primary method: sentence-transformer embeddings + cosine similarity (high accuracy).
    Fallback: TF-IDF vectorization + cosine similarity (if model failed to load).
    Returns a float between 0.0 (completely different) and 1.0 (identical)."""
    try:
        model = get_similarity_model()
        if model == "tfidf":
            # Model didn't load — use TF-IDF fallback
            return _calculate_similarity_tfidf(text1, text2)
        from sklearn.metrics.pairwise import cosine_similarity
        # Encode both texts into dense vector embeddings, then compute cosine similarity
        embeddings = model.encode([text1, text2])
        return float(cosine_similarity([embeddings[0]], [embeddings[1]])[0][0])
    except Exception:
        return _calculate_similarity_tfidf(text1, text2)

def _calculate_similarity_tfidf(text1: str, text2: str) -> float:
    """TF-IDF cosine similarity fallback. If sklearn is unavailable,
    uses Jaccard index (set intersection / set union) as a last resort."""
    try:
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.metrics.pairwise import cosine_similarity
        tfidf = TfidfVectorizer()
        matrix = tfidf.fit_transform([text1, text2])
        return float(cosine_similarity(matrix[0:1], matrix[1:2])[0][0])
    except Exception:
        # Jaccard similarity as ultimate fallback
        w1, w2 = set(text1.lower().split()), set(text2.lower().split())
        if not w1 or not w2:
            return 0.0
        return len(w1 & w2) / len(w1 | w2)

# ═══════════════════════════════════════════════════════════
# PDF PROCESSING
# ═══════════════════════════════════════════════════════════

def extract_identifiers_from_pdf_bytes(pdf_bytes: bytes) -> Dict[str, Optional[str]]:
    """Extract DOI, arXiv ID, and PMID from uploaded PDF bytes.
    Only scans the first 2 pages where identifiers typically appear."""
    identifiers: Dict[str, Optional[str]] = {"doi": None, "arxiv": None, "pmid": None}
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = ""
        for page_num in range(min(2, len(doc))):
            text += doc[page_num].get_text()
        doc.close()

        if not text.strip():
            return identifiers

        # DOI
        doi_match = re.search(DOI_PATTERN, text, re.IGNORECASE)
        if doi_match:
            doi = sanitize_doi(doi_match.group(0))
            if validate_doi(doi):
                identifiers["doi"] = doi

        # arXiv
        arxiv_match = re.search(ARXIV_PATTERN, text, re.IGNORECASE)
        if arxiv_match:
            arxiv_id = sanitize_arxiv_id(arxiv_match.group(1))
            if validate_arxiv_id(arxiv_id):
                identifiers["arxiv"] = arxiv_id

        # PMID
        pmid_match = re.search(PMID_PATTERN, text, re.IGNORECASE)
        if pmid_match:
            pmid = sanitize_pmid(pmid_match.group(1))
            if validate_pmid(pmid):
                identifiers["pmid"] = pmid
    except Exception:
        pass
    return identifiers


def extract_title_from_pdf_bytes(pdf_bytes: bytes) -> Optional[str]:
    """Heuristically extract the paper title from the first page of a PDF.
    Skips lines containing metadata keywords and picks the first plausible title line."""
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        if len(doc) == 0:
            return None
        text = doc[0].get_text()
        doc.close()
        if not text.strip():
            return None
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        skip_words = ["doi:", "http", "arxiv", "volume", "page", "journal",
                      "published", "copyright", "received", "accepted", "available", "pmid:", "pmc"]
        for line in lines[:15]:
            if any(s in line.lower() for s in skip_words):
                continue
            if 15 <= len(line) <= 300 and not line.isupper():
                return line
        if lines and len(lines[0]) >= 10:
            return lines[0]
        return None
    except Exception:
        return None


def extract_author_from_pdf_bytes(pdf_bytes: bytes) -> Optional[str]:
    """Heuristically extract author names from the first page of a PDF.
    Scores candidate lines based on capitalized words and separators (commas, 'and', '&')."""
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        if len(doc) == 0:
            return None
        text = doc[0].get_text()
        doc.close()
        if not text.strip():
            return None
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        potential = []
        skip_words = ["abstract", "introduction", "keywords", "doi:", "http",
                      "journal", "volume", "university", "department", "email",
                      "copyright", "rights reserved", "published", "manuscript",
                      "corresponding author", "affiliation", "received", "accepted",
                      "article", "citation", "license", "creative commons", "arxiv", "pmid:", "pmc"]
        for i, line in enumerate(lines[:60]):
            if len(line) < 3 or len(line) > 500:
                continue
            if any(s in line.lower() for s in skip_words):
                continue
            if re.search(r"[A-Z][a-z]+", line):
                words = line.split()
                caps = [w for w in words if w and len(w) > 1 and w[0].isupper()]
                if 1 <= len(caps) <= 25:
                    has_sep = "," in line or " and " in line.lower() or "&" in line
                    looks_names = sum(1 for w in caps if len(w) >= 2 and sum(c.isalpha() for c in w) >= len(w) * 0.7) >= len(caps) * 0.5
                    if has_sep or looks_names:
                        score = len(caps) + (10 if has_sep else 0)
                        potential.append((i, line, score))
        if potential:
            potential.sort(key=lambda x: x[2], reverse=True)
            return " ".join(potential[0][1].split())
        return None
    except Exception:
        return None

# ═══════════════════════════════════════════════════════════
# API FETCHERS (CACHED)
# ═══════════════════════════════════════════════════════════

@st.cache_data(show_spinner=False, ttl=3600)
def get_paper_metadata(query: str) -> Optional[Dict[str, Any]]:
    """Fetch metadata from Crossref via DOI or title search.
    Crossref is the largest DOI registration agency and provides metadata for
    journal articles, books, and conference papers. Uses the habanero library
    as a Python client for the Crossref REST API.
    Results are cached for 1 hour (ttl=3600) to avoid redundant API calls."""
    query = query.strip()
    is_doi = query.startswith("10.")
    if is_doi:
        query = sanitize_doi(query)
        if not validate_doi(query):
            return None
    else:
        ok, _ = validate_title(query)
        if not ok:
            return None

    def fetch():
        cr = Crossref()  # Initialize the Crossref API client
        if is_doi:
            # Direct DOI lookup — returns exact match
            result = cr.works(ids=query)
            return result["message"] if result and "message" in result else None
        else:
            # Title search — returns best match from Crossref's index
            results = cr.works(query=query, limit=1)
            if results and "message" in results and "items" in results["message"]:
                items = results["message"]["items"]
                return items[0] if items else None
            return None

    try:
        work = retry_with_backoff(fetch, max_attempts=MAX_RETRIES)
        if not work:
            return None
        # Extract and format author names from Crossref's structured author array
        # Each author object has 'given' (first name) and 'family' (last name) fields
        all_authors = work.get("author", [])
        if all_authors:
            author_string = ", ".join(
                f"{a.get('given', '')} {a.get('family', '')}".strip() for a in all_authors
            )
        else:
            author_string = "Unknown Author"
        # Build a standardized metadata dict with all fields needed for citation formatting
        return {
            "title": work.get("title", ["Unknown Title"])[0] if work.get("title") else "Unknown Title",
            "author": author_string,
            "year": str(work.get("published", {}).get("date-parts", [[None]])[0][0] or "n.d."),
            "doi": work.get("DOI", "Unknown DOI"),
            "abstract": clean_abstract(work.get("abstract", "")),
            "journal": html.unescape(work.get("container-title", [""])[0]) if work.get("container-title") else "",
            "volume": str(work.get("volume", "")) if work.get("volume") else "",
            "issue": str(work.get("issue", "")) if work.get("issue") else "",
            "pages": work.get("page", ""),
        }
    except Exception:
        return None


@st.cache_data(show_spinner=False, ttl=3600)
def get_metadata_from_arxiv(arxiv_id: str) -> Optional[Dict[str, Any]]:
    """Fetch metadata from the arXiv Atom API. Parses XML response to extract
    title, authors, year, DOI (if linked), and abstract."""
    arxiv_id = sanitize_arxiv_id(arxiv_id)
    if not validate_arxiv_id(arxiv_id):
        return None

    def fetch():
        resp = requests.get(ARXIV_API_URL, params={"id_list": arxiv_id, "max_results": 1}, timeout=10)
        if resp.status_code != 200:
            raise requests.RequestException(f"Status {resp.status_code}")
        root = ET.fromstring(resp.content)
        ns = {"atom": "http://www.w3.org/2005/Atom", "arxiv": "http://arxiv.org/schemas/atom"}
        entry = root.find("atom:entry", ns)
        if entry is None:
            return None
        title_el = entry.find("atom:title", ns)
        title = title_el.text.strip().replace("\n", " ") if title_el is not None else "Unknown Title"
        authors = []
        for a in entry.findall("atom:author", ns):
            n = a.find("atom:name", ns)
            if n is not None:
                authors.append(n.text.strip())
        pub_el = entry.find("atom:published", ns)
        year = pub_el.text[:4] if pub_el is not None else "n.d."
        doi_el = entry.find("arxiv:doi", ns)
        doi = doi_el.text.strip() if doi_el is not None else ""
        abs_el = entry.find("atom:summary", ns)
        abstract = abs_el.text.strip().replace("\n", " ") if abs_el is not None else ""
        return {"title": title, "author": ", ".join(authors) or "Unknown Author",
                "year": year, "doi": doi, "arxiv": arxiv_id, "abstract": clean_abstract(abstract)}

    try:
        return retry_with_backoff(fetch, max_attempts=MAX_RETRIES)
    except Exception:
        return None


def _fetch_pubmed_abstract(pmid: str) -> str:
    """Fetch the abstract from PubMed using the efetch API (XML).
    The esummary endpoint does NOT return abstracts, so this is required."""
    try:
        resp = requests.get(
            f"{PUBMED_API_URL}/efetch.fcgi",
            params={"db": "pubmed", "id": pmid, "rettype": "xml", "retmode": "xml"},
            timeout=10,
        )
        if resp.status_code != 200:
            return ""
        root = ET.fromstring(resp.content)
        # PubMed XML: <PubmedArticle><MedlineCitation><Article><Abstract><AbstractText>
        abstract_parts = []
        for abs_text in root.iter("AbstractText"):
            label = abs_text.get("Label", "")
            text = abs_text.text or ""
            # Also gather any tail text from child elements (e.g. <i>, <b> inside abstract)
            full_text = "".join(abs_text.itertext()).strip()
            if label and full_text:
                abstract_parts.append(f"{label}: {full_text}")
            elif full_text:
                abstract_parts.append(full_text)
        return " ".join(abstract_parts)
    except Exception:
        return ""


@st.cache_data(show_spinner=False, ttl=3600)
def get_metadata_from_pubmed(pmid: str) -> Optional[Dict[str, Any]]:
    """Fetch article metadata from PubMed's ESummary API (for bibliographic fields)
    and EFetch API (for the abstract). Returns title, authors, year, DOI,
    journal name, volume, issue, pages, and abstract."""
    pmid = sanitize_pmid(pmid)
    if not validate_pmid(pmid):
        return None

    def fetch():
        # Step 1: Get bibliographic metadata from esummary (fast, JSON)
        resp = requests.get(f"{PUBMED_API_URL}/esummary.fcgi",
                            params={"db": "pubmed", "id": pmid, "retmode": "json"}, timeout=10)
        if resp.status_code != 200:
            raise requests.RequestException(f"Status {resp.status_code}")
        data = resp.json()
        if "result" not in data or pmid not in data["result"]:
            return None
        r = data["result"][pmid]
        title = r.get("title", "Unknown Title").strip()
        authors = [a.get("name", "") for a in r.get("authors", [])]
        year = r.get("pubdate", "n.d.")[:4] if "pubdate" in r else "n.d."
        doi = ""
        if "elocationid" in r:
            eloc = r["elocationid"]
            if eloc.startswith("doi:"):
                doi = eloc[4:].strip()
            elif "10." in eloc:
                doi = eloc.strip()
        if not doi and "articleids" in r:
            for aid in r["articleids"]:
                if aid.get("idtype") == "doi":
                    doi = aid.get("value", "")
                    break
        journal = html.unescape(r.get("fulljournalname", "") or r.get("source", ""))
        volume = r.get("volume", "")
        issue = r.get("issue", "")
        pages = r.get("pages", "")

        # Step 2: Get abstract from efetch (esummary does NOT return abstracts)
        abstract = _fetch_pubmed_abstract(pmid)

        return {"title": title, "author": ", ".join(authors) or "Unknown Author",
                "year": year, "doi": doi, "pmid": pmid,
                "journal": journal, "volume": volume,
                "issue": issue, "pages": pages,
                "abstract": clean_abstract(abstract)}

    try:
        result = retry_with_backoff(fetch, max_attempts=MAX_RETRIES)
        if result:
            return result
    except Exception:
        pass

    # Fallback: if PubMed E-utilities failed (e.g. rate-limited on cloud hosting),
    # convert PMID to DOI via NCBI ID converter and fetch metadata from Crossref instead.
    try:
        conv_resp = requests.get(
            "https://www.ncbi.nlm.nih.gov/pmc/utils/idconv/v1.0/",
            params={"ids": pmid, "format": "json", "tool": "ScholarSync"},
            timeout=10,
        )
        if conv_resp.status_code == 200:
            conv_data = conv_resp.json()
            records = conv_data.get("records", [])
            if records:
                doi = records[0].get("doi", "")
                if doi and validate_doi(doi):
                    meta = get_paper_metadata(doi)
                    if meta:
                        meta["pmid"] = pmid  # Preserve the original PMID
                        return meta
    except Exception:
        pass

    return None


@st.cache_data(show_spinner=False, ttl=3600)
def get_metadata_from_isbn(isbn: str) -> Optional[Dict[str, Any]]:
    """Fetch book metadata from the Google Books API using an ISBN.
    The Google Books API is queried with 'isbn:<number>' as the search parameter.
    Returns title, authors, year, publisher, edition, and description.
    Results are cached for 1 hour to avoid redundant API calls."""
    isbn = sanitize_isbn(isbn)
    if not validate_isbn(isbn):
        return None

    def fetch():
        resp = requests.get(GOOGLE_BOOKS_API_URL, params={"q": f"isbn:{isbn}", "maxResults": 1}, timeout=10)
        if resp.status_code != 200:
            raise requests.RequestException(f"Status {resp.status_code}")
        data = resp.json()
        if "items" not in data or not data["items"]:
            return None
        info = data["items"][0]["volumeInfo"]
        pub_date = info.get("publishedDate", "n.d.")
        return {
            "title": info.get("title", "Unknown Title"),
            "author": ", ".join(info.get("authors", [])) or "Unknown Author",
            "year": pub_date[:4] if pub_date and len(pub_date) >= 4 else "n.d.",
            "publisher": info.get("publisher", ""),
            "isbn": isbn,
            "edition": info.get("printType", ""),
            "abstract": clean_abstract(info.get("description", "")),
            "type": "book",
        }

    try:
        return retry_with_backoff(fetch, max_attempts=MAX_RETRIES)
    except Exception:
        return None


@st.cache_data(show_spinner=False, ttl=3600)
def get_metadata_from_url(url: str) -> Optional[Dict[str, Any]]:
    """Scrape metadata from a web page (news article, blog, etc.) using BeautifulSoup.
    Extracts title, author, date, publication name, and description from HTML meta tags."""
    url = sanitize_url(url)
    if not validate_url(url):
        return None

    def fetch():
        # Use a browser-like User-Agent to avoid being blocked by websites
        headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
        resp = requests.get(url, headers=headers, timeout=10)
        if resp.status_code != 200:
            raise requests.RequestException(f"Status {resp.status_code}")
        soup = BeautifulSoup(resp.content, "html.parser")

        # Title
        title = None
        for getter in [
            lambda: (soup.find("meta", property="og:title") or {}).get("content"),
            lambda: (soup.find("meta", {"name": "twitter:title"}) or {}).get("content"),
            lambda: soup.title.string if soup.title else None,
            lambda: soup.find("h1").get_text().strip() if soup.find("h1") else None,
        ]:
            try:
                t = getter()
                if t:
                    title = t.strip()
                    break
            except Exception:
                continue
        title = title or "Unknown Title"

        # Author
        author = "Unknown Author"
        for getter in [
            lambda: (soup.find("meta", {"name": "author"}) or {}).get("content"),
            lambda: (soup.find("meta", property="article:author") or {}).get("content"),
            lambda: soup.find(["span", "div", "p"], class_=re.compile(r"author|byline", re.I)).get_text().strip(),
        ]:
            try:
                a = getter()
                if a and a.strip():
                    author = a.strip()
                    break
            except Exception:
                continue

        # Date
        year = "n.d."
        for getter in [
            lambda: (soup.find("meta", property="article:published_time") or {}).get("content"),
            lambda: (soup.find("meta", property="datePublished") or {}).get("content"),
            lambda: soup.find("time")["datetime"] if soup.find("time") else None,
        ]:
            try:
                d = getter()
                if d:
                    m = re.search(r"(\d{4})", d)
                    if m:
                        year = m.group(1)
                        break
            except Exception:
                continue

        # Publication name
        publication = "Unknown Publication"
        og_site = soup.find("meta", property="og:site_name")
        if og_site and og_site.get("content"):
            publication = og_site["content"].strip()
        else:
            publication = urlparse(url).netloc.replace("www.", "").split(".")[0].capitalize()

        # Description
        abstract = ""
        md = soup.find("meta", {"name": "description"})
        if md and md.get("content"):
            abstract = md["content"].strip()
        elif soup.find("meta", property="og:description"):
            abstract = soup.find("meta", property="og:description")["content"].strip()

        return {
            "title": title, "author": author, "year": year,
            "publication": publication, "url": url, "abstract": clean_abstract(abstract),
            "access_date": datetime.now().strftime("%B %d, %Y"), "type": "web-article",
        }

    try:
        return retry_with_backoff(fetch, max_attempts=MAX_RETRIES)
    except Exception:
        return None


@st.cache_data(show_spinner=False, ttl=3600)
def get_metadata_from_pmcid(pmcid: str) -> Optional[Dict[str, Any]]:
    """Fetch metadata for a PubMed Central article (PMCID).
    PMC doesn't have its own metadata API, so this function:
    1. Converts the PMCID to a PMID using NCBI's ID Converter API
    2. Fetches metadata via the PubMed (PMID) fetcher
    3. Falls back to DOI lookup via Crossref if PMID conversion fails."""
    pmcid = sanitize_pmcid(pmcid)
    if not validate_pmcid(pmcid):
        return None

    def fetch_pmid():
        resp = requests.get(
            "https://www.ncbi.nlm.nih.gov/pmc/utils/idconv/v1.0/",
            params={"ids": pmcid, "format": "json", "tool": "ScholarSync"},
            timeout=10,
        )
        if resp.status_code != 200:
            raise requests.RequestException(f"Status {resp.status_code}")
        data = resp.json()
        records = data.get("records", [])
        if not records:
            return None
        rec = records[0]
        # Return both pmid and doi so we can try either
        return {"pmid": rec.get("pmid", ""), "doi": rec.get("doi", "")}

    try:
        result = retry_with_backoff(fetch_pmid, max_attempts=MAX_RETRIES)
        if not result:
            return None
        # Try PMID first (gives full journal metadata)
        if result["pmid"] and validate_pmid(result["pmid"]):
            meta = get_metadata_from_pubmed(result["pmid"])
            if meta:
                return meta
        # Fallback to DOI via Crossref (also has full journal metadata)
        if result["doi"] and validate_doi(result["doi"]):
            meta = get_paper_metadata(result["doi"])
            if meta:
                return meta
        return None
    except Exception:
        return None


def get_metadata_with_fallback(identifiers: Dict[str, Optional[str]]) -> Optional[Dict[str, Any]]:
    """Try multiple APIs in priority order (DOI → arXiv → PMID) until one succeeds.
    Used when a PDF contains multiple identifiers."""
    if identifiers.get("doi"):
        meta = get_paper_metadata(identifiers["doi"])
        if meta:
            return meta
    if identifiers.get("arxiv"):
        meta = get_metadata_from_arxiv(identifiers["arxiv"])
        if meta:
            return meta
    if identifiers.get("pmid"):
        meta = get_metadata_from_pubmed(identifiers["pmid"])
        if meta:
            return meta
    return None

# ═══════════════════════════════════════════════════════════
# CITATION FORMATTING
# ═══════════════════════════════════════════════════════════

def format_author_names(author_string: str, style_name: str) -> str:
    """Reformat author names according to the chosen citation style.
    APA: Last, F. M.  |  MLA/Chicago: Last, First  |  Others: unchanged.
    Handles multiple authors with appropriate separators (& vs 'and')."""
    if not author_string or author_string == "Unknown Author":
        return author_string
    # Split the author string by common separators, trying the most specific first
    separators = [", and ", " and ", ", ", " & "]
    authors = [author_string]
    for sep in separators:
        if sep in authors[0]:
            authors = [a.strip() for a in authors[0].split(sep) if a.strip()]
            break
    formatted = []
    for author in authors:
        # Remove any stray superscript numbers (e.g. affiliation markers like "Smith1")
        author = re.sub(r"\d+", "", author).strip()
        parts = author.split()
        if not parts:
            continue
        if style_name == "APA":
            # APA format: Last, F. M.  (last name, then initials of remaining names)
            if len(parts) == 1:
                formatted.append(parts[0])
            else:
                initials = [f"{p[0]}." for p in parts[:-1] if p]
                formatted.append(f"{parts[-1]}, {' '.join(initials)}")
        elif style_name in ["MLA", "Chicago"]:
            # MLA/Chicago format: Last, First Middle
            if len(parts) == 1:
                formatted.append(parts[0])
            else:
                formatted.append(f"{parts[-1]}, {' '.join(parts[:-1])}")
        else:
            # IEEE, Harvard, Vancouver: keep author name as-is
            formatted.append(author)
    if not formatted:
        return author_string
    # Join multiple authors with style-appropriate separators
    if len(formatted) == 1:
        return formatted[0]
    if len(formatted) == 2:
        # APA uses "&", others use "and"
        return f"{formatted[0]}, & {formatted[1]}" if style_name == "APA" else f"{formatted[0]} and {formatted[1]}"
    if style_name == "APA":
        # APA 7th: list up to 20 authors; for 21+, show first 19 then "..." then last
        if len(formatted) <= 20:
            return ", ".join(formatted[:-1]) + f", & {formatted[-1]}"
        return ", ".join(formatted[:19]) + f", . . . {formatted[-1]}"
    return ", ".join(formatted[:-1]) + f", and {formatted[-1]}"


def get_parenthetical_citation(metadata: Dict[str, Any], style_name: str) -> str:
    """Generate an in-text/parenthetical citation, e.g. (Smith, 2023) or (Smith).
    Uses 'et al.' when multiple authors are present."""
    try:
        author_full = metadata.get("author", "Unknown")
        if "," in author_full:
            last_name = author_full.split(",")[0].strip()
        else:
            parts = author_full.split()
            last_name = parts[-1] if parts else "Unknown"
        if "," in author_full or " and " in author_full.lower():
            last_name = last_name.split()[0] + " et al."
        year = metadata.get("year", "n.d.")
        if style_name in ["APA", "Harvard", "Chicago"]:
            return f"({last_name}, {year})"
        elif style_name == "MLA":
            return f"({last_name})"
        return f"({last_name}, {year})"
    except Exception:
        return "(Citation Error)"


def format_citation_from_metadata(metadata: Dict[str, Any], style_name: str) -> str:
    """Route metadata to the appropriate formatter based on source type.
    Each source type (journal article, book, web article) has different
    required fields and formatting rules. Defaults to journal article
    formatting if no type is specified."""
    try:
        source_type = metadata.get("type", "journal")
        if source_type == "book":
            return _format_book_citation(metadata, style_name)
        elif source_type == "web-article":
            return _format_web_article_citation(metadata, style_name)
        return _format_journal_citation(metadata, style_name)
    except Exception:
        return f"{metadata.get('author', 'Unknown')} ({metadata.get('year', 'n.d.')}). {metadata.get('title', 'Unknown Title')}."


def _format_journal_citation(m: Dict[str, Any], style: str) -> str:
    """Format a journal article citation in the specified style (APA, MLA, etc.).
    Each style has distinct rules for punctuation, ordering, and emphasis:
    - APA: Author (Year). Title. Journal, Volume(Issue), Pages. DOI-URL
    - MLA: Author. \"Title.\" Journal, vol. V, no. N, Year, pp. Pages. doi:DOI
    - Harvard: Author (Year) 'Title', Journal, Volume(Issue), pp. Pages. doi: DOI
    - Chicago: Author. \"Title.\" Journal Volume, no. Issue (Year): Pages. DOI-URL
    - IEEE: Author, \"Title,\" Journal, vol. V, no. N, pp. Pages, Year. doi: DOI
    - Vancouver: Author. Title. Journal. Year;Volume(Issue):Pages. doi: DOI"""
    # Extract all metadata fields, using empty strings as defaults
    fa = format_author_names(m.get("author", "Unknown Author"), style)
    t = m.get("title", "Unknown Title")
    y = m.get("year", "n.d.")
    doi = m.get("doi", "")
    j = m.get("journal", "")
    v = m.get("volume", "")
    iss = m.get("issue", "")
    p = m.get("pages", "")
    has_doi = doi and doi != "Unknown DOI" and validate_doi(doi)

    # Wrap journal name in *italic* markers for styles that require it.
    # APA: journal + volume italic; MLA/Chicago/IEEE: journal italic; Harvard: journal + volume italic
    # Vancouver: journal NOT italic (abbreviated journal names stay plain)
    if style == "APA":
        c = f"{fa} ({y}). {t}."
        if j:
            c += f" *{j}*"
            if v:
                c += f", *{v}*"
                if iss:
                    c += f"({iss})"
            if p:
                c += f", {p}"
            c += "."
        if has_doi:
            c += f" https://doi.org/{doi}"
    elif style == "MLA":
        c = f'{fa}. "{t}."'
        if j:
            c += f" *{j}*"
            if v: c += f", vol. {v}"
            if iss: c += f", no. {iss}"
            c += f", {y}"
            if p: c += f", pp. {p}"
            c += "."
        else:
            c += f" {y}."
        if has_doi:
            c += f" doi:{doi}"
    elif style == "Harvard":
        c = f"{fa} ({y}) '{t}',"
        if j:
            c += f" *{j}*"
            if v:
                c += f", *{v}*"
                if iss: c += f"({iss})"
            if p: c += f", pp. {p}"
        c += "."
        if has_doi:
            c += f" doi: {doi}"
    elif style == "Chicago":
        c = f'{fa}. "{t}."'
        if j:
            c += f" *{j}*"
            if v: c += f" {v}"
            if iss: c += f", no. {iss}"
            c += f" ({y})"
            if p: c += f": {p}"
            c += "."
        else:
            c += f" {y}."
        if has_doi:
            c += f" https://doi.org/{doi}."
    elif style == "IEEE":
        c = f'{fa}, "{t},"'
        if j:
            c += f" *{j}*"
            if v: c += f", vol. {v}"
            if iss: c += f", no. {iss}"
            if p: c += f", pp. {p}"
        c += f", {y}."
        if has_doi:
            c += f" doi: {doi}."
    elif style == "Vancouver":
        c = f"{fa}. {t}."
        if j:
            c += f" {j}. {y}"
            if v:
                c += f";{v}"
                if iss: c += f"({iss})"
            if p: c += f":{p}"
            c += "."
        else:
            c += f" {y}."
        if has_doi:
            c += f" doi: {doi}"
    else:
        c = f"{fa} ({y}). {t}."
        if has_doi:
            c += f" https://doi.org/{doi}"
    return c


def _format_book_citation(m: Dict[str, Any], style: str) -> str:
    """Format a book citation with publisher and edition info.
    Book titles are italicized in APA, MLA, Harvard, and Chicago styles."""
    fa = format_author_names(m.get("author", "Unknown Author"), style)
    t = m.get("title", "Unknown Title")
    y = m.get("year", "n.d.")
    pub = m.get("publisher", "")
    ed = m.get("edition", "")
    if style == "APA":
        c = f"{fa} ({y}). *{t}*"
        if ed: c += f" ({ed})"
        c += "."
        if pub: c += f" {pub}."
    elif style == "MLA":
        c = f"{fa}. *{t}*."
        if ed: c += f" {ed},"
        if pub: c += f" {pub},"
        c += f" {y}."
    elif style == "Harvard":
        c = f"{fa} ({y}) *{t}*."
        if ed: c += f" {ed}."
        if pub: c += f" {pub}."
    elif style == "Chicago":
        c = f"{fa}. {y}. *{t}*."
        if ed: c += f" {ed}."
        if pub: c += f" {pub}."
    elif style == "IEEE":
        c = f"{fa}, *{t}*"
        if ed: c += f", {ed}"
        c += "."
        if pub: c += f" {pub},"
        c += f" {y}."
    elif style == "Vancouver":
        c = f"{fa}. {t}."
        if ed: c += f" {ed}."
        if pub: c += f" {pub};"
        c += f" {y}."
    else:
        c = f"{fa} ({y}). *{t}*."
        if pub: c += f" {pub}."
    return c


def _format_web_article_citation(m: Dict[str, Any], style: str) -> str:
    """Format a web/news article citation with URL and access date.
    Publication/website name is italicized in APA, MLA, and Chicago styles."""
    fa = format_author_names(m.get("author", "Unknown Author"), style)
    t = m.get("title", "Unknown Title")
    y = m.get("year", "n.d.")
    pub = m.get("publication", "")
    url = m.get("url", "")
    access = m.get("access_date", datetime.now().strftime("%B %d, %Y"))
    # Wrap publication name in *italic* markers for APA, MLA, and Chicago
    # Harvard, IEEE, and Vancouver do NOT italicize website names
    if style == "APA":
        c = f"{fa} ({y}). {t}."
        if pub: c += f" *{pub}*."
        if url: c += f" {url}"
    elif style == "MLA":
        c = f'{fa}. "{t}."'
        if pub: c += f" *{pub}*,"
        c += f" {y},"
        if url: c += f" {url}."
        c += f" Accessed {access}."
    elif style == "Harvard":
        c = f"{fa} ({y}) '{t}'"
        if pub: c += f", {pub}"
        c += "."
        if url: c += f" Available at: {url} (Accessed: {access})."
    elif style == "Chicago":
        c = f'{fa}. "{t}."'
        if pub: c += f" *{pub}*,"
        c += f" {y}."
        if url: c += f" {url}."
    elif style == "IEEE":
        c = f'{fa}, "{t},"'
        if pub: c += f" {pub},"
        c += f" {y}."
        if url: c += f" [Online]. Available: {url}"
    elif style == "Vancouver":
        c = f"{fa}. {t}."
        if pub: c += f" {pub}."
        c += f" {y}."
        if url: c += f" Available from: {url}"
    else:
        c = f"{fa} ({y}). {t}."
        if pub: c += f" *{pub}*."
        if url: c += f" {url}"
    return c

# ═══════════════════════════════════════════════════════════
# DUPLICATE DETECTION
# ═══════════════════════════════════════════════════════════

def check_duplicate_reference(
    new_ref: Dict[str, Any],
    existing_refs: List[Dict[str, Any]],
    title_threshold: float = 0.85,
    author_year_threshold: float = 0.70,
) -> Tuple[bool, Optional[Dict[str, Any]], Optional[str]]:
    """Check if a new reference duplicates one already in the collection.
    Uses three strategies: (1) exact identifier match, (2) title similarity
    via ML embeddings, (3) author + year similarity. Returns (is_dup, existing_ref, reason)."""
    if not existing_refs:
        return False, None, None
    nm = new_ref.get("meta", {})
    for er in existing_refs:
        em = er.get("meta", {})
        # Strategy 1: Exact identifier (DOI/arXiv/PMID) match
        for key in ["doi", "arxiv", "pmid"]:
            nv, ev = nm.get(key, "").strip(), em.get(key, "").strip()
            if nv and ev and nv != "Unknown DOI" and nv.lower() == ev.lower():
                return True, er, f"Identical {key.upper()}: {nv}"
        # Strategy 2: Semantic title similarity (ML cosine similarity >= 85%)
        nt, et = nm.get("title", "").strip(), em.get("title", "").strip()
        if nt and et and nt != "Unknown Title":
            sim = calculate_similarity_advanced(nt, et)
            if sim >= title_threshold:
                return True, er, f"Similar title ({sim:.0%} match)"
        # Strategy 3: Same author names + same publication year
        na, ea = nm.get("author", "").lower(), em.get("author", "").lower()
        ny, ey = str(nm.get("year", "")), str(em.get("year", ""))
        if na and ea and ny and ey and ny == ey:
            if calculate_similarity_advanced(na, ea) >= author_year_threshold:
                return True, er, f"Same author & year"
    return False, None, None

# ═══════════════════════════════════════════════════════════
# AUTO-DETECTION & PIPELINE
# ═══════════════════════════════════════════════════════════

def _extract_identifier_from_url(url: str) -> Optional[Tuple[str, str]]:
    """Try to extract a known identifier (DOI, arXiv, PMID, PMCID, ISBN) from a URL.
    Returns (type, value) or None if the URL is not a recognised identifier link."""
    url_lower = url.lower()

    # DOI links: https://doi.org/10.xxxx/... or https://dx.doi.org/10.xxxx/...
    if "doi.org/" in url_lower:
        doi = sanitize_doi(url)
        if validate_doi(doi):
            return "doi", doi

    # Some publisher URLs embed the DOI in the path (e.g. /doi/10.xxxx/...)
    doi_in_path = re.search(r"/(10\.\d{4,9}/[^\s?#]+)", url)
    if doi_in_path:
        doi = doi_in_path.group(1).rstrip("/")
        if validate_doi(doi):
            return "doi", doi

    # arXiv links: https://arxiv.org/abs/2301.12345  or /pdf/2301.12345
    arxiv_match = re.search(r"arxiv\.org/(?:abs|pdf)/(\d{4}\.\d{4,5}(?:v\d+)?)", url_lower)
    if arxiv_match:
        aid = arxiv_match.group(1)
        if validate_arxiv_id(aid):
            return "arxiv", aid

    # PubMed links: https://pubmed.ncbi.nlm.nih.gov/12345678/
    pmid_match = re.search(r"pubmed\.ncbi\.nlm\.nih\.gov/(\d{1,9})", url_lower)
    if pmid_match:
        pid = pmid_match.group(1)
        if validate_pmid(pid):
            return "pmid", pid

    # PMC links: https://www.ncbi.nlm.nih.gov/pmc/articles/PMC1234567/
    #        or: https://pmc.ncbi.nlm.nih.gov/articles/PMC11312261/
    pmc_match = re.search(r"(?:pmc\.ncbi\.nlm\.nih\.gov|ncbi\.nlm\.nih\.gov/pmc)/articles/(PMC\d{1,9})", url, re.IGNORECASE)
    if pmc_match:
        pcid = pmc_match.group(1).upper()
        if validate_pmcid(pcid):
            return "pmcid", pcid

    # Google Books links containing ISBN
    isbn_match = re.search(r"[?&](?:isbn|vid=ISBN)[:=](\d{10,13})", url, re.IGNORECASE)
    if isbn_match:
        isbn = isbn_match.group(1)
        if validate_isbn(isbn):
            return "isbn", isbn

    return None


def detect_identifier_type(identifier_input: str) -> Tuple[str, str]:
    """Auto-detect the type of identifier from user input.
    Priority: URL → DOI → ISBN → arXiv → PMID → PMCID → Title (fallback).
    Returns (type_string, cleaned_value)."""
    if not identifier_input or not isinstance(identifier_input, str):
        return "invalid", ""
    identifier_input = identifier_input.strip()
    # If the input looks like a URL, first try to extract a known identifier from it
    if identifier_input.startswith(("http://", "https://", "www.")):
        extracted = _extract_identifier_from_url(identifier_input)
        if extracted:
            return extracted
        # No known identifier found – treat as a generic URL
        url = sanitize_url(identifier_input)
        if validate_url(url):
            return "url", url
    # DOI
    if identifier_input.startswith("10.") or "doi.org/" in identifier_input.lower():
        doi = sanitize_doi(identifier_input)
        if validate_doi(doi):
            return "doi", doi
    # ISBN
    if "ISBN" in identifier_input.upper() or re.match(
        r"^(97[89])?\d{9}[\dX]$", identifier_input.replace("-", "").replace(" ", "")
    ):
        isbn = sanitize_isbn(identifier_input)
        if validate_isbn(isbn):
            return "isbn", isbn
    # arXiv
    if "arxiv" in identifier_input.lower() or re.match(r"^\d{4}\.\d{4,5}(v\d+)?$", identifier_input):
        aid = sanitize_arxiv_id(identifier_input)
        if validate_arxiv_id(aid):
            return "arxiv", aid
    # PMID
    if identifier_input.upper().startswith("PMID:") or (identifier_input.isdigit() and 1 <= len(identifier_input) <= 9):
        pid = sanitize_pmid(identifier_input)
        if validate_pmid(pid):
            return "pmid", pid
    # PMCID
    if "PMC" in identifier_input.upper():
        pcid = sanitize_pmcid(identifier_input)
        if validate_pmcid(pcid):
            return "pmcid", pcid
    # Title fallback
    ok, _ = validate_title(identifier_input)
    if ok:
        return "title", identifier_input
    return "invalid", identifier_input


def process_identifier(identifier_input: str, style_name: str = "APA") -> Optional[Dict[str, Any]]:
    """Core pipeline: detect type → fetch metadata → format citation.
    This is the main entry point for the Universal Search tab.
    Steps:
    1. Auto-detect what kind of identifier the user entered (DOI, ISBN, arXiv, etc.)
    2. Call the appropriate API fetcher to retrieve metadata
    3. Format both a full citation and a parenthetical/in-text citation
    4. Return everything as a dict for display and collection storage."""
    id_type, id_value = detect_identifier_type(identifier_input)
    if id_type == "invalid":
        return None

    metadata = None
    # Dispatch to the correct API fetcher based on detected identifier type
    if id_type == "doi":
        metadata = get_paper_metadata(id_value)          # Crossref API
    elif id_type == "arxiv":
        metadata = get_metadata_from_arxiv(id_value)     # arXiv Atom API
    elif id_type == "pmid":
        metadata = get_metadata_from_pubmed(id_value)    # PubMed E-Utilities
    elif id_type == "pmcid":
        metadata = get_metadata_from_pmcid(id_value)     # NCBI ID Converter → PubMed
    elif id_type == "isbn":
        metadata = get_metadata_from_isbn(id_value)      # Google Books API
    elif id_type == "url":
        metadata = get_metadata_from_url(id_value)       # HTML scraping
    elif id_type == "title":
        metadata = get_paper_metadata(id_value)           # Crossref title search

    if not metadata:
        return None

    # Package metadata + formatted citations into a result dict
    return {
        "meta": metadata,                                           # Raw metadata from API
        "full": format_citation_from_metadata(metadata, style_name),  # Full bibliography entry
        "parenthetical": get_parenthetical_citation(metadata, style_name),  # In-text citation
        "style": style_name,                                        # Citation style used
        "detected_type": id_type,                                   # What type was auto-detected
    }


def process_pdf(pdf_bytes: bytes, style_name: str = "APA") -> Optional[Dict[str, Any]]:
    """Process uploaded PDF: first try to find identifiers (DOI/arXiv/PMID) in the text
    and fetch metadata via APIs. If no identifiers found, fall back to heuristic
    title and author extraction from the PDF text itself."""
    identifiers = extract_identifiers_from_pdf_bytes(pdf_bytes)
    metadata = None
    if any(identifiers.values()):
        metadata = get_metadata_with_fallback(identifiers)
    if not metadata:
        # Fallback: extract title & author directly from PDF text
        title = extract_title_from_pdf_bytes(pdf_bytes)
        author = extract_author_from_pdf_bytes(pdf_bytes)
        if title:
            metadata = {
                "title": title,
                "author": author or "Unknown Author",
                "year": "n.d.",
                "doi": "",
            }
    if not metadata:
        return None
    return {
        "meta": metadata,
        "full": format_citation_from_metadata(metadata, style_name),
        "parenthetical": get_parenthetical_citation(metadata, style_name),
        "style": style_name,
        "detected_type": "pdf",
    }

# ═══════════════════════════════════════════════════════════
# WORD EXPORT (IN-MEMORY)
# ═══════════════════════════════════════════════════════════

def export_references_to_word(
    references: List[Dict[str, Any]],
    title: str = "References",
    sort_by_author: bool = True,
    include_abstracts: bool = True,
) -> Optional[bytes]:
    """Build a Word (.docx) document entirely in memory and return the raw bytes.
    The document is structured into three sections:
      1. Bibliography — full formatted citations with hanging indentation
      2. In-Text Citations — parenthetical citations for use within paper body
      3. Abstracts — paper abstracts (if available and include_abstracts=True)
    Uses python-docx to construct the document. The bytes can be downloaded
    directly via Streamlit's download button without writing to disk."""
    if not references:
        return None
    try:
        doc = Document()
        # Set 1-inch margins on all sides (standard academic formatting)
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        h = doc.add_heading(title, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        info = doc.add_paragraph()
        info.add_run(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n").italic = True
        info.add_run(f"Total References: {len(references)}\n").italic = True
        info.add_run(f"Citation Style: {references[0]['style']}\n").italic = True
        info.add_run("Generated by ScholarSync").italic = True
        doc.add_paragraph()
        doc.add_paragraph("─" * 80)
        doc.add_paragraph()

        refs = references.copy()
        if sort_by_author:
            # Sort references alphabetically by the first author's last name
            # (standard in APA, MLA, Harvard, and Chicago bibliographies)
            def _sort_key(x):
                author = x["meta"].get("author", "")
                if not author or author == "Unknown Author":
                    return "zzz"
                # Get first author (before any comma separating multiple authors)
                first_author = author.split(",")[0].strip()
                # Last name is the last word of "FirstName LastName"
                parts = first_author.split()
                return parts[-1].lower() if parts else "zzz"
            refs.sort(key=_sort_key)

        # --- Section 1: Bibliography ---
        # Each citation uses a hanging indent (first line left, rest indented 0.5")
        # which is the standard format for reference lists
        bib_h = doc.add_heading("Bibliography", level=2)
        bib_h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph()

        def _add_run(para, text, italic=False):
            """Helper: add a text run to a Word paragraph with consistent font styling."""
            if not text:
                return
            run = para.add_run(text)
            run.font.size = Pt(11)
            run.font.name = "Times New Roman"
            if italic:
                run.italic = True

        def _add_citation(para, citation_text, meta, style):
            """Add citation text to a Word paragraph with proper italic formatting
            for journal names, volume numbers, book titles, and publication/website names."""
            citation_text = citation_text.replace("*", "")
            journal = meta.get("journal", "")
            publication = meta.get("publication", "")
            volume = meta.get("volume", "")
            title = meta.get("title", "")
            source_type = meta.get("type", "journal")

            # Determine which name to italicize based on source type and style
            italic_name = ""
            if journal and journal in citation_text:
                italic_name = journal
            elif publication and publication in citation_text and style in ["APA", "MLA", "Chicago"]:
                italic_name = publication
            elif source_type == "book" and title and title in citation_text and style != "Vancouver":
                italic_name = title

            if italic_name:
                idx = citation_text.find(italic_name)
                _add_run(para, citation_text[:idx])
                _add_run(para, italic_name, italic=True)
                after = citation_text[idx + len(italic_name):]
                if volume and style in ["APA", "Harvard"] and journal:
                    vol_pat = rf"(,\s*)({re.escape(volume)})(\(?)"
                    vm = re.search(vol_pat, after)
                    if vm:
                        _add_run(para, after[:vm.start()])
                        _add_run(para, vm.group(1))
                        _add_run(para, vm.group(2), italic=True)
                        _add_run(para, after[vm.end() - len(vm.group(3)):])
                    else:
                        _add_run(para, after)
                else:
                    _add_run(para, after)
            else:
                _add_run(para, citation_text)

        for ref in refs:
            # Create paragraph with hanging indent (first line at 0", continuation at 0.5")
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.space_after = Pt(12)
            _add_citation(p, ref["full"], ref.get("meta", {}), ref.get("style", "APA"))

        doc.add_page_break()

        # --- Section 2: In-Text Citations ---
        it_h = doc.add_heading("In-Text Citations", level=2)
        it_h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph()
        exp = doc.add_paragraph()
        exp_run = exp.add_run("Use these in-text citations when referencing the papers in your work:")
        exp_run.font.size = Pt(10)
        exp_run.italic = True
        doc.add_paragraph()

        for i, ref in enumerate(refs, 1):
            itp = doc.add_paragraph()
            nr = itp.add_run(f"[{i}] ")
            nr.bold = True
            nr.font.size = Pt(11)
            nr.font.name = "Times New Roman"
            cr = itp.add_run(ref["parenthetical"])
            cr.font.size = Pt(11)
            cr.font.name = "Times New Roman"
            tp = doc.add_paragraph()
            tp.paragraph_format.left_indent = Inches(0.3)
            ttxt = f"→ {ref['meta']['title'][:80]}{'...' if len(ref['meta']['title']) > 80 else ''}"
            tr = tp.add_run(ttxt)
            tr.font.size = Pt(9)
            tr.italic = True
            tp.paragraph_format.space_after = Pt(12)

        # --- Section 3: Abstracts ---
        if include_abstracts:
            abs_refs = [r for r in refs if r["meta"].get("abstract")]
            if abs_refs:
                doc.add_page_break()
                ah = doc.add_heading("Abstracts", level=2)
                ah.alignment = WD_ALIGN_PARAGRAPH.LEFT
                doc.add_paragraph()
                ae = doc.add_paragraph()
                ae_run = ae.add_run("Abstracts extracted from paper metadata (where available):")
                ae_run.font.size = Pt(10)
                ae_run.italic = True
                doc.add_paragraph()
                for idx, ref in enumerate(abs_refs, 1):
                    hp = doc.add_paragraph()
                    hnr = hp.add_run(f"[{idx}] ")
                    hnr.bold = True; hnr.font.size = Pt(11); hnr.font.name = "Times New Roman"
                    hcr = hp.add_run(f"{ref['meta']['author']} ({ref['meta']['year']})")
                    hcr.font.size = Pt(11); hcr.font.name = "Times New Roman"; hcr.bold = True
                    ttp = doc.add_paragraph()
                    ttp.paragraph_format.left_indent = Inches(0.3)
                    ttr = ttp.add_run(ref["meta"]["title"])
                    ttr.font.size = Pt(10); ttr.italic = True; ttr.font.name = "Times New Roman"
                    abp = doc.add_paragraph()
                    abp.paragraph_format.left_indent = Inches(0.3)
                    abp.paragraph_format.right_indent = Inches(0.3)
                    abp.paragraph_format.space_after = Pt(16)
                    abs_text = ref["meta"]["abstract"]
                    if len(abs_text) > 1500:
                        abs_text = abs_text[:1500] + "... [truncated]"
                    abr = abp.add_run(abs_text)
                    abr.font.size = Pt(10); abr.font.name = "Times New Roman"
                    doc.add_paragraph("─" * 60)
                    doc.add_paragraph()

        # Save document to an in-memory buffer (no temp files needed)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception:
        return None

# ═══════════════════════════════════════════════════════════
# ABSTRACT PROCESSING
# ═══════════════════════════════════════════════════════════

def process_abstract(abstract: str) -> Dict[str, Any]:
    """Process an abstract for display: compute word count and truncate if over 150 words."""
    wc = len(abstract.split()) if abstract else 0
    display = (" ".join(abstract.split()[:150]) + "...") if wc > 150 else abstract
    return {"full": abstract, "display": display, "word_count": wc}

# ═══════════════════════════════════════════════════════════
# HELPER: COPYABLE CITATION WITH ITALIC FORMATTING
# ═══════════════════════════════════════════════════════════

def _render_copyable_citation(citation_text: str, unique_key: str):
    """Render a citation with italic formatting and a copy button that preserves
    italic formatting when pasted into Word, Google Docs, etc.
    Uses navigator.clipboard.write() with text/html MIME type for rich text copy.
    Falls back to plain text copy if rich text API is unavailable."""
    import html as _html

    # Convert *text* markers to <i>text</i> HTML tags
    parts = citation_text.split("*")
    html_content = ""
    for i, part in enumerate(parts):
        escaped = _html.escape(part)
        if i % 2 == 1:  # Odd-indexed parts are between * markers → italic
            html_content += f"<i>{escaped}</i>"
        else:
            html_content += escaped

    # Estimate component height based on text length (accounts for line wrapping)
    estimated_lines = max(1, len(citation_text) // 70 + 1)
    height = 30 + estimated_lines * 28

    # Build the HTML component with dark/light mode support and copy functionality
    component_html = f"""
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{ background: transparent; }}
        .citation-wrapper {{
            position: relative;
            padding: 12px 44px 12px 14px;
            border-radius: 8px;
            font-family: 'Source Serif Pro', 'Georgia', serif;
            font-size: 15px;
            line-height: 1.65;
            border: 1px solid;
        }}
        /* Light mode (default) */
        .citation-wrapper {{
            background-color: #f0f2f6;
            border-color: #d1d5db;
            color: #1a1a2e;
        }}
        .copy-btn {{
            background: #e2e4e8;
            color: #555;
            border: 1px solid #c5c8cd;
        }}
        .copy-btn:hover {{ background: #d1d4d9; }}
        /* Dark mode */
        @media (prefers-color-scheme: dark) {{
            .citation-wrapper {{
                background-color: #262730;
                border-color: #4a4b57;
                color: #e8e8ed;
            }}
            .copy-btn {{
                background: #3b3c4a;
                color: #ccc;
                border: 1px solid #555;
            }}
            .copy-btn:hover {{ background: #4a4b5a; }}
        }}
        .copy-btn {{
            position: absolute;
            top: 8px;
            right: 8px;
            cursor: pointer;
            border-radius: 5px;
            padding: 3px 7px;
            font-size: 13px;
            transition: background 0.2s;
        }}
    </style>
    <div class="citation-wrapper" id="cite-{unique_key}">
        <span id="text-{unique_key}">{html_content}</span>
        <button class="copy-btn" id="btn-{unique_key}" onclick="copyRich_{unique_key.replace('-','_')}()"
                title="Copy with italics">📋</button>
    </div>
    <script>
        async function copyRich_{unique_key.replace('-','_')}() {{
            const textEl = document.getElementById('text-{unique_key}');
            const btn = document.getElementById('btn-{unique_key}');
            const richHtml = textEl.innerHTML;
            const plainText = textEl.innerText;
            try {{
                const htmlBlob = new Blob([richHtml], {{type: 'text/html'}});
                const textBlob = new Blob([plainText], {{type: 'text/plain'}});
                await navigator.clipboard.write([
                    new ClipboardItem({{'text/html': htmlBlob, 'text/plain': textBlob}})
                ]);
                btn.textContent = '✅';
                setTimeout(() => btn.textContent = '📋', 1500);
            }} catch(e) {{
                try {{
                    await navigator.clipboard.writeText(plainText);
                    btn.textContent = '✅';
                    setTimeout(() => btn.textContent = '📋', 1500);
                }} catch(e2) {{
                    btn.textContent = '❌';
                    setTimeout(() => btn.textContent = '📋', 1500);
                }}
            }}
        }}
    </script>
    """
    components.html(component_html, height=height)


# ═══════════════════════════════════════════════════════════
# HELPER: DISPLAY RESULT CARD
# ═══════════════════════════════════════════════════════════

def _show_result(result: Dict[str, Any], key_prefix: str = ""):
    """Display a citation result card in the Streamlit UI.
    Shows metadata fields, abstract (collapsible), formatted citations,
    and an 'Add to Collection' button that checks for duplicates first."""
    meta = result["meta"]
    detected = result.get("detected_type", "unknown").upper()

    st.markdown(f"**Detected type:** `{detected}`")

    col1, col2 = st.columns([3, 1])  # 3:1 ratio — main info on left, volume/issue/pages on right
    with col1:
        st.markdown(f"**Title:** {meta.get('title', 'N/A')}")
        st.markdown(f"**Author:** {meta.get('author', 'N/A')}")
        st.markdown(f"**Year:** {meta.get('year', 'N/A')}")
        for field, label in [("doi", "DOI"), ("arxiv", "arXiv"), ("pmid", "PMID"),
                             ("isbn", "ISBN"), ("publisher", "Publisher"),
                             ("journal", "Journal"), ("url", "URL")]:
            val = meta.get(field)
            if val:
                st.markdown(f"**{label}:** {val}")
    with col2:
        vol = meta.get("volume", "")
        iss = meta.get("issue", "")
        pgs = meta.get("pages", "")
        if vol:
            st.markdown(f"**Volume:** {vol}")
        if iss:
            st.markdown(f"**Issue:** {iss}")
        if pgs:
            st.markdown(f"**Pages:** {pgs}")

    # Abstract
    if meta.get("abstract"):
        with st.expander("📄 Abstract", expanded=False):
            ab = process_abstract(meta["abstract"])
            st.write(ab["full"])
            st.caption(f"{ab['word_count']} words")

    # Citations — rendered with italics + copy button that preserves formatting.
    # Uses a custom HTML component via streamlit.components for rich-text clipboard copy.
    st.markdown("---")
    st.markdown("**Full Citation:**")
    _render_copyable_citation(result["full"], f"full-{key_prefix}")
    st.markdown("**In-Text Citation:**")
    _render_copyable_citation(result["parenthetical"], f"intext-{key_prefix}")

    # "Add to Collection" button with ML-based duplicate checking
    btn_key = f"add_{key_prefix}_{meta.get('title', '')[:30]}"
    if st.button("➕ Add to Collection", key=btn_key, type="primary"):
        # If duplicate detection is enabled, check before adding
        if st.session_state.get("duplicate_detection", True):
            is_dup, dup_ref, reason = check_duplicate_reference(result, st.session_state.collection)
            if is_dup:
                st.session_state.pending_duplicate = {"result": result, "reason": reason, "dup": dup_ref, "key": btn_key}
                st.rerun()
        else:
            is_dup = False
        if not st.session_state.get("pending_duplicate"):
            st.session_state.collection.append(result)
            st.toast("✅ Added to collection!")
            st.rerun()

# ═══════════════════════════════════════════════════════════
# SIDEBAR
# The sidebar persists across all tabs and contains:
# - Citation style selector (APA, MLA, Harvard, Chicago, IEEE, Vancouver)
# - Reference counter and Word export download button
# - Duplicate detection toggle (ML-powered)
# ═══════════════════════════════════════════════════════════

with st.sidebar:
    st.image("https://img.icons8.com/fluency/96/book-shelf.png", width=60)
    st.title("ScholarSync")
    st.caption("Smart Reference Generator")
    st.divider()

    # Citation style dropdown — maps display names to internal style keys
    style_display = st.selectbox("📝 Citation Style", list(STYLE_DISPLAY_NAMES.keys()), index=0)
    style = STYLE_DISPLAY_NAMES[style_display]  # e.g. "APA 7th Edition" → "APA"
    st.divider()

    st.metric("📚 References Collected", len(st.session_state.collection))

    if st.session_state.collection:
        word_bytes = export_references_to_word(st.session_state.collection, sort_by_author=True)
        if word_bytes:
            st.download_button(
                "⬇️ Download Word (.docx)",
                data=word_bytes,
                file_name=f"ScholarSync_References_{datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

        if st.button("🗑️ Clear All References", use_container_width=True):
            st.session_state.collection.clear()
            st.session_state.last_result = None
            st.session_state.pending_duplicate = None
            st.rerun()

    st.divider()
    st.toggle("🔍 Duplicate Detection", value=True, key="duplicate_detection",
             help="Uses an ML similarity model to detect duplicate references before adding.")
    st.divider()
    st.caption("Supports: DOI · arXiv · PMID · ISBN · URL · PDF · Title")

# ═══════════════════════════════════════════════════════════
# DUPLICATE WARNING MODAL
# When a duplicate is detected, this section shows a warning banner
# with the option to "Add Anyway" (override) or "Cancel".
# Uses st.session_state.pending_duplicate to persist across reruns.
# ═══════════════════════════════════════════════════════════

if st.session_state.pending_duplicate:
    dup_info = st.session_state.pending_duplicate
    st.warning(f"⚠️ **Duplicate detected:** {dup_info['reason']}")
    dup_meta = dup_info["dup"]["meta"]
    st.caption(f"Existing: *{dup_meta.get('title', 'N/A')[:80]}* — {dup_meta.get('author', 'N/A')[:50]} ({dup_meta.get('year', 'N/A')})")
    c1, c2 = st.columns(2)
    with c1:
        if st.button("Add Anyway", type="primary"):
            st.session_state.collection.append(dup_info["result"])
            st.session_state.pending_duplicate = None
            st.toast("✅ Added (duplicate override)")
            st.rerun()
    with c2:
        if st.button("Cancel"):
            st.session_state.pending_duplicate = None
            st.rerun()
    st.divider()

# ═══════════════════════════════════════════════════════════
# MAIN AREA
# The app uses 4 tabs for different input methods:
#   1. Universal Search — auto-detects identifier type and fetches metadata
#   2. PDF Upload — extracts identifiers from uploaded PDF files
#   3. Manual Entry — form-based input for any source type
#   4. My Collection — view, manage, and analyze saved references
# ═══════════════════════════════════════════════════════════

st.title("📚 ScholarSync")
st.caption("Paste a DOI, ISBN, arXiv ID, PMID, URL, title, or upload a PDF to generate citations instantly.")

tab_search, tab_pdf, tab_manual, tab_collection = st.tabs(
    ["🔍 Universal Search", "📄 PDF Upload", "✍️ Manual Entry", "📚 My Collection"]
)

# ── Tab 1: Universal Search ──────────────────────────────
# Accepts any identifier type: DOI, ISBN, arXiv, PMID, PMCID, URL, or title.
# The detect_identifier_type() function auto-classifies the input,
# then process_identifier() fetches metadata from the appropriate API.

with tab_search:
    identifier = st.text_input(
        "Paste a DOI, ISBN, arXiv ID, PMID, URL, or Title",
        placeholder="e.g. 10.1038/nature12345  or  https://www.bbc.com/news/article  or  978-0-7475-3269-9",
        key="search_input",
    )
    if st.button("Generate Citation", key="search_go", type="primary", use_container_width=True):
        if not identifier.strip():
            st.warning("Please enter an identifier or title.")
        else:
            with st.spinner("Fetching metadata..."):
                result = process_identifier(identifier.strip(), style)
            if result:
                st.session_state.last_result = result
            else:
                st.error("❌ Could not retrieve metadata. Check the identifier and try again.")

    if st.session_state.last_result and st.session_state.last_result.get("detected_type") != "pdf_tab":
        _show_result(st.session_state.last_result, key_prefix="search")

# ── Tab 2: PDF Upload ─────────────────────────────────────
# Users can upload a research paper PDF. The app scans the first 2 pages
# for identifiers (DOI, arXiv, PMID), then fetches metadata via APIs.
# If no identifiers are found, it falls back to heuristic title/author extraction.

with tab_pdf:
    uploaded = st.file_uploader("Upload a research paper PDF", type=["pdf"], key="pdf_upload")
    if uploaded is not None:
        pdf_bytes = uploaded.read()
        with st.spinner("Extracting identifiers from PDF..."):
            result = process_pdf(pdf_bytes, style)
        if result:
            result["detected_type"] = "pdf_tab"
            _show_result(result, key_prefix="pdf")
        else:
            st.error("❌ Could not extract enough information from the PDF.")

# ── Tab 3: Manual Entry ───────────────────────────────────
# A form-based interface for manually entering citation details.
# Supports three source types, each with different required fields:
# - Journal Article: title, author, year, DOI, journal, volume, issue, pages
# - Book: title, author, year, ISBN, publisher, edition
# - Web/News Article: title, author, year, URL, publication name, access date

with tab_manual:
    m_source_type = st.radio(
        "Source Type",
        ["Journal Article", "Book", "Web / News Article"],
        horizontal=True,
        key="manual_source_type",
    )
    with st.form("manual_form"):
        m_title = st.text_input("Title *", placeholder="The Impact of Machine Learning on Healthcare")
        m_author = st.text_input("Author(s) *", placeholder="John Smith, Jane Doe")
        mc1, mc2 = st.columns(2)
        with mc1:
            m_year = st.text_input("Year", value="n.d.", placeholder="2023")
        with mc2:
            if m_source_type == "Journal Article":
                m_doi = st.text_input("DOI (optional)", placeholder="10.1234/example")
            elif m_source_type == "Book":
                m_isbn = st.text_input("ISBN (optional)", placeholder="978-0-7475-3269-9")
            else:
                m_url = st.text_input("URL *", placeholder="https://www.bbc.com/news/article")

        if m_source_type == "Journal Article":
            mc3, mc4 = st.columns(2)
            with mc3:
                m_journal = st.text_input("Journal (optional)", placeholder="Nature")
            with mc4:
                m_volume = st.text_input("Volume (optional)", placeholder="46")
            mc5, mc6 = st.columns(2)
            with mc5:
                m_issue = st.text_input("Issue (optional)", placeholder="2")
            with mc6:
                m_pages = st.text_input("Pages (optional)", placeholder="123-145")
        elif m_source_type == "Book":
            mc3, mc4 = st.columns(2)
            with mc3:
                m_publisher = st.text_input("Publisher (optional)", placeholder="Oxford University Press")
            with mc4:
                m_edition = st.text_input("Edition (optional)", placeholder="3rd")
        else:  # Web / News Article
            mc3, mc4 = st.columns(2)
            with mc3:
                m_publication = st.text_input("Publication / Website Name", placeholder="BBC News")
            with mc4:
                m_access_date = st.text_input("Access Date", value=datetime.now().strftime("%B %d, %Y"))

        submitted = st.form_submit_button("Generate Citation", type="primary", use_container_width=True)
        if submitted:
            if not m_title.strip() or not m_author.strip():
                st.warning("Title and Author are required.")
            elif m_source_type == "Web / News Article" and not m_url.strip():
                st.warning("URL is required for web/news articles.")
            else:
                if m_source_type == "Journal Article":
                    meta = {
                        "title": m_title.strip(),
                        "author": m_author.strip(),
                        "year": m_year.strip() or "n.d.",
                        "doi": sanitize_doi(m_doi.strip()),  # Strip doi:/URL prefixes
                        "journal": m_journal.strip(),
                        "volume": m_volume.strip(),
                        "issue": m_issue.strip(),
                        "pages": m_pages.strip(),
                    }
                elif m_source_type == "Book":
                    meta = {
                        "title": m_title.strip(),
                        "author": m_author.strip(),
                        "year": m_year.strip() or "n.d.",
                        "publisher": m_publisher.strip(),
                        "edition": m_edition.strip(),
                        "isbn": m_isbn.strip() if m_isbn.strip() else "",
                        "type": "book",
                    }
                else:  # Web / News Article
                    meta = {
                        "title": m_title.strip(),
                        "author": m_author.strip(),
                        "year": m_year.strip() or "n.d.",
                        "publication": m_publication.strip(),
                        "url": m_url.strip(),
                        "access_date": m_access_date.strip(),
                        "type": "web-article",
                    }
                st.session_state.manual_result = {
                    "meta": meta,
                    "full": format_citation_from_metadata(meta, style),
                    "parenthetical": get_parenthetical_citation(meta, style),
                    "style": style,
                    "detected_type": "manual",
                }

    # Display result outside the form so st.button works
    if st.session_state.get("manual_result"):
        _show_result(st.session_state.manual_result, key_prefix="manual")

# ── Tab 4: Collection & Analytics ─────────────────────────
# Displays all saved references with:
# - A bar chart showing publication year distribution (built with Counter)
# - Summary statistics (total count, earliest/latest year, year range)
# - Expandable cards for each reference with full citation, in-text citation,
#   abstract preview, and a remove button

with tab_collection:
    refs = st.session_state.collection
    if not refs:
        st.info("No references collected yet. Use the other tabs to generate and add citations.")
    else:
        st.subheader(f"📚 {len(refs)} Reference(s)")

        # Year analytics
        years = []
        for ref in refs:
            ys = str(ref["meta"].get("year", "")).strip()
            if ys.isdigit() and 1900 <= int(ys) <= 2100:
                years.append(int(ys))

        if years:
            counts = Counter(years)
            chart_data = {str(y): counts[y] for y in sorted(counts)}
            st.bar_chart(chart_data)
            scol1, scol2, scol3, scol4 = st.columns(4)
            scol1.metric("Total", len(years))
            scol2.metric("Earliest", min(years))
            scol3.metric("Latest", max(years))
            scol4.metric("Range", f"{max(years) - min(years)} yrs")

        st.divider()

        # List each reference
        for i, ref in enumerate(refs):
            with st.expander(f"**[{i+1}]** {ref['meta'].get('title', 'N/A')[:80]}", expanded=False):
                st.markdown(f"**Author:** {ref['meta'].get('author', 'N/A')}")
                st.markdown(f"**Year:** {ref['meta'].get('year', 'N/A')}")
                _render_copyable_citation(ref["full"], f"col-full-{i}")
                _render_copyable_citation(ref["parenthetical"], f"col-intext-{i}")
                if ref["meta"].get("abstract"):
                    ab = process_abstract(ref["meta"]["abstract"])
                    st.caption(f"Abstract ({ab['word_count']} words): {ab['display'][:200]}...")
                if st.button("🗑️ Remove", key=f"remove_{i}"):
                    st.session_state.collection.pop(i)
                    st.rerun()
